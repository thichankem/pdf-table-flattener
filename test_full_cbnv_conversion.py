import os
import sys
import docx
from doc_table_converter import convert_doc_or_pdf_to_docx, clean_table_grid, format_table_to_dash_text, clean_text_string, replace_table_inplace, is_header_footer_table, is_bullet_marker, unpack_text_table_inplace, fix_toc_paragraphs, cleanup_document_blank_spaces

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def process_nested_tables_in_cell(cell, separator=" | ", use_header=True, show_row_indices=False, bullet_prefix=True):
    if not hasattr(cell, 'tables') or not cell.tables:
        return
    for nested_table in list(cell.tables):
        # Process nested tables recursively first
        for nr in nested_table.rows:
            for nc in nr.cells:
                process_nested_tables_in_cell(nc, separator, use_header, show_row_indices, bullet_prefix)
        
        ngrid = clean_table_grid(nested_table)
        if is_header_footer_table(ngrid):
            tbl_elm = nested_table._element
            if tbl_elm.getparent() is not None:
                tbl_elm.getparent().remove(tbl_elm)
            continue
            
        formatted_text = format_table_to_dash_text(ngrid, separator=separator, use_header=use_header, show_row_indices=show_row_indices, bullet_prefix=bullet_prefix)
        
        tbl_elm = nested_table._element
        parent_elm = tbl_elm.getparent()
        tbl_index = parent_elm.index(tbl_elm)
        
        if formatted_text:
            p = cell.add_paragraph()
            p.text = formatted_text
            parent_elm.insert(tbl_index, p._element)
        parent_elm.remove(tbl_elm)

pdf_path = "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf"
real_docx_path, is_temp = convert_doc_or_pdf_to_docx(pdf_path)

doc = docx.Document(real_docx_path)

# 1. Process nested tables first
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            process_nested_tables_in_cell(c)

# 2. Fix TOC
fix_toc_paragraphs(doc)

# 3. Merge cross-page continuation tables
raw_tables = list(doc.tables)
table_groups = []

for table in raw_tables:
    grid = clean_table_grid(table)
    if not grid:
        tbl_elm = table._element
        if tbl_elm.getparent() is not None:
            tbl_elm.getparent().remove(tbl_elm)
        continue
    
    num_cols = max(len(r) for r in grid)
    
    is_continuation = False
    if table_groups and num_cols >= 3:
        prev_group = table_groups[-1]
        prev_grid = prev_group[-1][1]
        prev_cols = max(len(r) for r in prev_grid) if prev_grid else 0
        
        if prev_cols == num_cols and len(grid) >= 1:
            first_row = grid[0]
            col0 = clean_text_string(first_row[0]) if len(first_row) > 0 else ""
            col1 = clean_text_string(first_row[1]) if len(first_row) > 1 else ""
            col2 = clean_text_string(first_row[2]) if len(first_row) > 2 else ""
            
            if not col0 and not col1 and col2:
                is_continuation = True
    
    if is_continuation:
        table_groups[-1].append((table, grid))
    else:
        table_groups.append([(table, grid)])

extracted_texts = []
for group in table_groups:
    if len(group) == 1:
        table, grid = group[0]
    else:
        first_table = group[0][0]
        merged_grid = list(group[0][1])
        
        for tbl_idx in range(1, len(group)):
            cont_table, cont_grid = group[tbl_idx]
            
            if cont_grid:
                first_cont_row = cont_grid[0]
                col2_content = clean_text_string(first_cont_row[2]) if len(first_cont_row) > 2 else ""
                
                if col2_content and merged_grid:
                    last_row = merged_grid[-1]
                    if len(last_row) > 2:
                        existing = clean_text_string(last_row[2])
                        if existing:
                            last_row[2] = existing + "\n" + col2_content
                        else:
                            last_row[2] = col2_content
                
                for r_idx in range(1, len(cont_grid)):
                    merged_grid.append(cont_grid[r_idx])
            
            tbl_elm = cont_table._element
            if tbl_elm.getparent() is not None:
                tbl_elm.getparent().remove(tbl_elm)
        
        table = first_table
        grid = merged_grid

    if is_header_footer_table(grid):
        tbl_elm = table._element
        if tbl_elm.getparent() is not None:
            tbl_elm.getparent().remove(tbl_elm)
        continue
        
    num_rows = len(grid)
    num_cols = max(len(r) for r in grid) if num_rows > 0 else 0

    if num_cols == 1 or (num_rows == 1 and is_bullet_marker(grid[0][0])):
        unpack_text_table_inplace(doc, table, grid)
        continue

    table_text = format_table_to_dash_text(grid, separator=" | ", use_header=True, show_row_indices=False, bullet_prefix=True)
    if table_text:
        replace_table_inplace(doc, table, table_text)
        extracted_texts.append(table_text)

cleanup_document_blank_spaces(doc)

out_pdf = "output/Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx_convert.pdf"
out_docx = "temp_cbnv_test.docx"
doc.save(out_docx)

from doc_table_converter import convert_docx_to_pdf
convert_docx_to_pdf(out_docx, out_pdf)

if os.path.exists(out_docx):
    os.remove(out_docx)
if is_temp and os.path.exists(real_docx_path):
    os.remove(real_docx_path)

print("✅ Conversion finished!")
