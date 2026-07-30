import sys
import os
import docx
from doc_table_converter import convert_doc_or_pdf_to_docx, clean_table_grid, format_table_to_dash_text, clean_text_string, replace_table_inplace, is_header_footer_table, is_bullet_marker, fix_toc_paragraphs, cleanup_document_blank_spaces
from docx.shared import Pt

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def unpack_table_paragraphs_inplace(doc, table):
    tbl_elm = table._element
    parent_elm = tbl_elm.getparent()
    tbl_index = parent_elm.index(tbl_elm)

    inserted_count = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                txt = clean_text_string(p.text)
                if txt:
                    np = doc.add_paragraph()
                    np.text = txt
                    np.paragraph_format.space_before = Pt(0)
                    np.paragraph_format.space_after = Pt(2)
                    parent_elm.insert(tbl_index + inserted_count, np._element)
                    inserted_count += 1
    
    parent_elm.remove(tbl_elm)

def process_nested_tables_in_cell(cell):
    if not hasattr(cell, 'tables') or not cell.tables:
        return
    for nested_table in list(cell.tables):
        for nr in nested_table.rows:
            for nc in nr.cells:
                process_nested_tables_in_cell(nc)
        
        ngrid = clean_table_grid(nested_table)
        if is_header_footer_table(ngrid):
            tbl_elm = nested_table._element
            if tbl_elm.getparent() is not None:
                tbl_elm.getparent().remove(tbl_elm)
            continue
            
        formatted_text = format_table_to_dash_text(ngrid)
        
        tbl_elm = nested_table._element
        parent_elm = tbl_elm.getparent()
        tbl_index = parent_elm.index(tbl_elm)
        
        if formatted_text:
            lines = formatted_text.splitlines()
            for l_idx, line_str in enumerate(lines):
                p = cell.add_paragraph()
                p.text = line_str
                parent_elm.insert(tbl_index + l_idx, p._element)
        parent_elm.remove(nested_table._element)

pdf_path = "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf"
real_docx_path, is_temp = convert_doc_or_pdf_to_docx(pdf_path)
doc = docx.Document(real_docx_path)

# 1. Process nested tables
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            process_nested_tables_in_cell(c)

# 2. Check Table 2
t2 = doc.tables[1]
print("Table 2 before unpack:")
print(f"  Rows: {len(t2.rows)}, Cells per row: {[len(r.cells) for r in t2.rows]}")
for r in t2.rows:
    for c in r.cells:
        print("  Cell text:", repr(c.text))

print("\nUnpacking Table 2...")
unpack_table_paragraphs_inplace(doc, t2)

print("\nDocument paragraphs count:", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs[:30]):
    print(f"P{i}: {p.text[:100]}")

if is_temp and os.path.exists(real_docx_path):
    os.remove(real_docx_path)
