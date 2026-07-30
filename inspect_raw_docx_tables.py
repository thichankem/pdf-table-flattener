import os
import sys
import docx
from doc_table_converter import convert_doc_or_pdf_to_docx, clean_text_string

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

pdf_path = "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf"
docx_path, is_temp = convert_doc_or_pdf_to_docx(pdf_path)

print(f"Docx path: {docx_path}")
doc = docx.Document(docx_path)

print(f"Total tables in docx: {len(doc.tables)}")

for t_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {t_idx + 1} (rows: {len(table.rows)}) ---")
    for r_idx, row in enumerate(table.rows):
        cells_str = []
        for c_idx, cell in enumerate(row.cells):
            t = clean_text_string(cell.text)
            if len(t) > 60:
                t = t[:57] + "..."
            cells_str.append(f"C{c_idx}: '{t}'")
        print(f"  R{r_idx}: " + " | ".join(cells_str))

        # Check for nested tables inside cells!
        for c_idx, cell in enumerate(row.cells):
            if cell.tables:
                print(f"    ⚠️ NESTED TABLE inside R{r_idx} C{c_idx}! Count: {len(cell.tables)}")
                for nt_idx, nt in enumerate(cell.tables):
                    print(f"      Nested table {nt_idx+1}: {len(nt.rows)} rows")
                    for nr_idx, nr in enumerate(nt.rows):
                        nt_cells = [clean_text_string(c.text) for c in nr.cells]
                        print(f"        NR{nr_idx}: {' | '.join(nt_cells)}")

if is_temp and os.path.exists(docx_path):
    try:
        os.remove(docx_path)
    except:
        pass
