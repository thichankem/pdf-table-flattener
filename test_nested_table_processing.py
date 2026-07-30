import sys
import os
import docx
from doc_table_converter import convert_doc_or_pdf_to_docx, clean_text_string, format_table_to_dash_text, clean_table_grid

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def process_nested_tables_in_cell(cell, separator=" | ", use_header=True, show_row_indices=False, bullet_prefix=True):
    if not cell.tables:
        return
    for nested_table in list(cell.tables):
        # Process nested tables recursively if any
        for nr in nested_table.rows:
            for nc in nr.cells:
                process_nested_tables_in_cell(nc, separator, use_header, show_row_indices, bullet_prefix)
        
        # Format the nested table
        ngrid = clean_table_grid(nested_table)
        formatted_text = format_table_to_dash_text(ngrid, separator=separator, use_header=use_header, show_row_indices=show_row_indices, bullet_prefix=bullet_prefix)
        
        # Replace the nested table in cell with formatted_text
        tbl_elm = nested_table._element
        parent_elm = tbl_elm.getparent()
        tbl_index = parent_elm.index(tbl_elm)
        
        # Create new paragraph with formatted_text
        p = cell.add_paragraph()
        p.text = formatted_text
        parent_elm.insert(tbl_index, p._element)
        parent_elm.remove(tbl_elm)

pdf_path = "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf"
docx_path, is_temp = convert_doc_or_pdf_to_docx(pdf_path)

doc = docx.Document(docx_path)

# First pass: process all nested tables inside cells of top-level tables
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            process_nested_tables_in_cell(c)

print("--- AFTER NESTED TABLES PROCESSED ---")
for t_idx, table in enumerate(doc.tables):
    grid = clean_table_grid(table)
    print(f"\nTABLE {t_idx+1} (rows: {len(grid)}, cols: {max(len(r) for r in grid) if grid else 0}):")
    for r_idx, r in enumerate(grid):
        print(f"  R{r_idx}: {' | '.join(r)}")

if is_temp and os.path.exists(docx_path):
    try:
        os.remove(docx_path)
    except:
        pass
