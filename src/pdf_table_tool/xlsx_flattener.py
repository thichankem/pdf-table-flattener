"""Flatten every table of an Excel workbook into bullet paragraphs.

A spreadsheet is the easy case of the three: it *is* a grid, so there is no
detection step at all -- no ruling lines to find, no whitespace gutters to
guess.  The only real work is deciding where one table ends and the next
begins, and turning a stored value back into the text the reader sees in Excel
(``0.291`` under a percent format is "29.1%", not "0.291").

The grid is then handed to the same :class:`TableFormatter` the PDF and Word
paths use, so identical tables produce identical bullets in all three formats.

Unlike a PDF or a .docx, a workbook has no in-place flattened form -- a bullet
list is not a grid of cells -- so the output is written as a new Word document,
one heading per sheet.
"""

import datetime
import logging
import re
import zipfile
from typing import Any, Dict, List, Optional, Sequence, Set, Tuple

from docx import Document
from docx.shared import Pt

from .config import settings
from .docx_flattener import (
    COL_WIDTH,
    INDENT_STEP_PT,
    ROW_HEIGHT,
    _cell_from_items,
    declared_header_structure,
)
from .formatter import INDENT_UNIT, TableFormatter, detect_structure, normalise_sliced_cells
from .grid_extractor import Grid, Item
from .outline import DocumentOutline, number_table_lines
from .text_utils import normalize_text

logger = logging.getLogger(__name__)

# (first_row, last_row, first_col, last_col), 1-based and inclusive, as openpyxl
# addresses cells.
Block = Tuple[int, int, int, int]

_DECIMALS_RE = re.compile(r"\.(0+)")

# One ``<c>`` element from the point its formula starts, and the cached result
# such a cell carries when Excel has actually calculated it.  A writer that only
# stores the formula leaves the value out (openpyxl) or writes it empty
# (``<v></v>``), and the number is then simply not in the file.
_FORMULA_CELL_RE = re.compile(rb"<f[ />].*?</c>", re.DOTALL)
_CACHED_VALUE_RE = re.compile(rb"<v[^>]*>\s*[^<\s]")


# ---------------------------------------------------------------------------
# reading a cell the way Excel shows it
# ---------------------------------------------------------------------------
def _decimals(fmt: str) -> int:
    match = _DECIMALS_RE.search(fmt)
    return len(match.group(1)) if match else 0


def _number_text(value: float, fmt: str) -> str:
    """Render a stored number through its own display format.

    Excel keeps 29.1% as 0.291 and a price as a bare float; the number format is
    the only record of what the sheet actually shows.  Ignoring it would put
    figures in the output that appear nowhere in the source.
    """
    # A format string carries up to four sections (positive;negative;zero;text).
    fmt = (fmt or "").split(";")[0]
    percent = "%" in fmt
    if percent:
        value = value * 100
    grouped = "#,#" in fmt or "0,0" in fmt
    places = _decimals(fmt)

    if places:
        body = f"{value:,.{places}f}" if grouped else f"{value:.{places}f}"
    elif float(value).is_integer():
        body = f"{int(value):,}" if grouped else str(int(value))
    else:
        # "General": show what the author typed, without float noise.
        body = f"{value:,}" if grouped else f"{value:g}"
    return f"{body}%" if percent else body


def _cell_text(cell) -> str:
    """The text a reader sees in one cell (empty string when there is none)."""
    value = cell.value
    if value is None:
        return ""
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    if isinstance(value, datetime.datetime):
        if value.time() == datetime.time(0, 0):
            return value.strftime("%d/%m/%Y")
        return value.strftime("%d/%m/%Y %H:%M")
    if isinstance(value, datetime.date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, datetime.time):
        return value.strftime("%H:%M")
    if isinstance(value, (int, float)):
        return _number_text(float(value), cell.number_format or "")
    return normalize_text(str(value))


# ---------------------------------------------------------------------------
# finding the tables on a sheet
# ---------------------------------------------------------------------------
def _merge_map(worksheet) -> Tuple[Dict[Tuple[int, int], Tuple[int, int]], Set[Tuple[int, int]]]:
    """``({anchor: (row_span, col_span)}, {positions the anchor swallowed})``."""
    anchors: Dict[Tuple[int, int], Tuple[int, int]] = {}
    covered: Set[Tuple[int, int]] = set()
    for rng in worksheet.merged_cells.ranges:
        anchor = (rng.min_row, rng.min_col)
        anchors[anchor] = (
            rng.max_row - rng.min_row + 1,
            rng.max_col - rng.min_col + 1,
        )
        for row in range(rng.min_row, rng.max_row + 1):
            for col in range(rng.min_col, rng.max_col + 1):
                if (row, col) != anchor:
                    covered.add((row, col))
    return anchors, covered


def sheet_blocks(worksheet, anchors=None, covered=None) -> List[Block]:
    """Split a sheet into tables: a fully blank row separates two of them.

    A merged cell counts as filling every row it covers, so a title merged down
    the side of a table does not cut the table in half.
    """
    if anchors is None or covered is None:
        anchors, covered = _merge_map(worksheet)

    max_row, max_col = worksheet.max_row or 0, worksheet.max_column or 0
    spans: List[Optional[Tuple[int, int]]] = []
    for row in range(1, max_row + 1):
        cols = [
            col
            for col in range(1, max_col + 1)
            if (row, col) not in covered
            and _cell_text(worksheet.cell(row=row, column=col)).strip()
        ]
        spans.append((min(cols), max(cols)) if cols else None)

    # Widen every row a non-empty merged cell reaches into, so the run of rows
    # below its anchor stays part of the same block.
    for (row, col), (row_span, col_span) in anchors.items():
        if not _cell_text(worksheet.cell(row=row, column=col)).strip():
            continue
        for r in range(row, min(row + row_span, max_row + 1)):
            lo, hi = col, col + col_span - 1
            current = spans[r - 1]
            spans[r - 1] = (lo, hi) if current is None else (
                min(current[0], lo), max(current[1], hi)
            )

    blocks: List[Block] = []
    run: List[Tuple[int, Tuple[int, int]]] = []
    for index, span in enumerate(spans, start=1):
        if span is None:
            if run:
                blocks.append(_block_of(run))
                run = []
        else:
            run.append((index, span))
    if run:
        blocks.append(_block_of(run))
    return blocks


def _block_of(run: Sequence[Tuple[int, Tuple[int, int]]]) -> Block:
    return (
        run[0][0],
        run[-1][0],
        min(span[0] for _row, span in run),
        max(span[1] for _row, span in run),
    )


_PRINT_TITLE_RE = re.compile(r"\$?(\d+):\$?(\d+)")


def declared_header_rows(worksheet, block: Block) -> int:
    """Leading rows of `block` the author set as Excel's repeating print titles.

    "Rows to repeat at top" is the spreadsheet's way of saying "these rows are
    the header", exactly as Word's "Repeat as header row" is.  It only counts
    when it starts where the block does -- print titles are set per sheet, and
    the second table on a sheet is not covered by them.
    """
    titles = worksheet.print_title_rows
    if not titles:
        return 0
    match = _PRINT_TITLE_RE.search(str(titles))
    if match is None:
        return 0
    first, last = int(match.group(1)), int(match.group(2))
    if first != block[0]:
        return 0
    return max(0, min(last, block[1]) - first + 1)


def build_grid_from_block(worksheet, block: Block, anchors, covered) -> Grid:
    """Build a :class:`Grid` for one table of a sheet.

    The synthetic geometry comes from the Word path, which spaces rows wider
    apart than the cell-slicing tolerance -- a spreadsheet row is always a real
    row, and nothing here should ever be fused with its neighbour.
    """
    first_row, last_row, first_col, last_col = block
    n_rows = last_row - first_row + 1
    n_cols = last_col - first_col + 1

    cells = []
    for row in range(first_row, last_row + 1):
        for col in range(first_col, last_col + 1):
            if (row, col) in covered:
                continue
            row_span, col_span = anchors.get((row, col), (1, 1))
            # A merge may reach past the block; the part outside is another table.
            row_span = max(1, min(row_span, last_row - row + 1))
            col_span = max(1, min(col_span, last_col - col + 1))
            text = _cell_text(worksheet.cell(row=row, column=col))
            items = [
                Item(level=0, text=line.strip())
                for line in text.split("\n")
                if line.strip()
            ]
            cells.append(
                _cell_from_items(
                    row - first_row, col - first_col, row_span, col_span, items
                )
            )

    cells.sort(key=lambda c: (c.row, c.col))
    return Grid(
        n_rows=n_rows,
        n_cols=n_cols,
        cells=cells,
        bbox=(0.0, 0.0, n_cols * COL_WIDTH, n_rows * ROW_HEIGHT),
    )


def has_uncached_formulas(path: str) -> bool:
    """Does the workbook hold a formula whose result was never stored?

    A cell reads as empty in that case, and it reads as empty on *both* sides of
    the verifier -- so the run would report "nothing lost" while a whole column
    is missing.  This is the one failure mode the token check cannot see, which
    is why it is looked for directly.
    """
    try:
        with zipfile.ZipFile(path) as archive:
            for name in archive.namelist():
                if not (
                    name.startswith("xl/worksheets/") and name.endswith(".xml")
                ):
                    continue
                blob = archive.read(name)
                for match in _FORMULA_CELL_RE.finditer(blob):
                    if not _CACHED_VALUE_RE.search(match.group(0)):
                        return True
    except (zipfile.BadZipFile, KeyError, OSError) as exc:  # pragma: no cover
        logger.debug("Formula scan skipped: %s", exc)
    return False


def workbook_texts(path: str) -> List[str]:
    """Every cell's visible text, for the verifier to check nothing was lost."""
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True)
    try:
        texts: List[str] = []
        for worksheet in workbook.worksheets:
            _anchors, covered = _merge_map(worksheet)
            for row in range(1, (worksheet.max_row or 0) + 1):
                for col in range(1, (worksheet.max_column or 0) + 1):
                    if (row, col) in covered:
                        continue
                    text = _cell_text(worksheet.cell(row=row, column=col))
                    if text.strip():
                        texts.append(text)
        return texts
    finally:
        workbook.close()


# ---------------------------------------------------------------------------
# writing the bullets out
# ---------------------------------------------------------------------------
def _add_bullet(document, line: str, size: Optional[float]) -> None:
    indent = len(line) - len(line.lstrip(" "))
    level = indent // max(1, len(INDENT_UNIT))
    paragraph = document.add_paragraph()
    fmt = paragraph.paragraph_format
    if level:
        fmt.left_indent = Pt(INDENT_STEP_PT * level)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(2)
    run = paragraph.add_run(line.strip())
    if size:
        run.font.size = Pt(size)


def write_docx(
    sheets: List[Tuple[str, List[str]]],
    output_path: str,
    label_sheets: Optional[bool] = None,
) -> None:
    """One heading per sheet, then one paragraph per bullet line."""
    document = Document()
    # A single-sheet workbook needs no heading; naming it "Sheet1" would add a
    # word to the document that appears nowhere in the data.
    if label_sheets is None:
        label_sheets = len([s for s in sheets if s[1]]) > 1
    for name, lines in sheets:
        if not lines:
            continue
        if label_sheets:
            document.add_heading(name, level=2)
        for line in lines:
            if line.strip():
                _add_bullet(document, line, settings.BULLET_FONT_SIZE)
    if not document.paragraphs:
        document.add_paragraph()
    document.save(output_path)


# ---------------------------------------------------------------------------
# pipeline
# ---------------------------------------------------------------------------
class XlsxTableFlattener:
    """Same contract as :class:`PDFTableFlattenerPipeline`, for .xlsx files."""

    def __init__(self, verify_output: bool = True, numbering: bool = True):
        self.formatter = TableFormatter()
        self.verify_output = verify_output
        self.numbering = numbering

    def sheet_lines(self, worksheet) -> List[Tuple[List[str], List[str], bool]]:
        """One entry per block of a sheet: ``(bullets, headers, is a table)``.

        The blocks stay apart because each table of a sheet is numbered on its
        own, and because a title merged across the top is a block that is *not*
        a table -- counting it would inflate the report and numbering it would
        promote a caption to a section.
        """
        anchors, covered = _merge_map(worksheet)
        blocks: List[Tuple[List[str], List[str], bool]] = []
        for block in sheet_blocks(worksheet, anchors, covered):
            grid = build_grid_from_block(worksheet, block, anchors, covered)
            if grid.n_rows == 0:
                continue

            # Excel states what a page cannot: a header merged over sub-labels
            # ("Doanh thu" above "Trước thuế" and "Tỷ lệ"), and which rows the
            # author repeats at the top of every printed page.  Both go through
            # the same pass the Word path uses.
            grid, structure = declared_header_structure(
                grid, declared_header_rows(worksheet, block)
            )
            normalised = normalise_sliced_cells(grid)
            if structure is None:
                structure = detect_structure(normalised)

            block_lines, headers = self.formatter.format_grid(grid, None, structure)
            if block_lines:
                blocks.append(
                    (block_lines, headers, grid.n_rows >= 2 and grid.n_cols >= 2)
                )
        return blocks

    def flatten_sheet(self, worksheet) -> Tuple[List[str], int]:
        """``(bullet_lines, tables_flattened)`` for one worksheet, unnumbered."""
        blocks = self.sheet_lines(worksheet)
        lines = [line for block, _headers, _is_table in blocks for line in block]
        return lines, sum(1 for _b, _h, is_table in blocks if is_table)

    def process(self, input_path: str, output_path: str) -> Dict[str, Any]:
        from openpyxl import load_workbook

        logger.info("Processing %s", input_path)
        # `data_only` reads the value Excel last calculated; a formula string
        # would put "=SUM(B2:B9)" in the output instead of the number on screen.
        workbook = load_workbook(input_path, data_only=True)
        try:
            read = [(ws.title, self.sheet_lines(ws)) for ws in workbook.worksheets]
        finally:
            workbook.close()

        # A sheet is the outline a workbook has: it names its tables and nothing
        # else does.  It only becomes a section when it is written out as a
        # heading -- numbering against a heading the reader cannot see would
        # point at nothing.
        label_sheets = len([1 for _t, blocks in read if blocks]) > 1
        outline = DocumentOutline()

        sheets: List[Tuple[str, List[str]]] = []
        total_tables = 0
        all_lines: List[str] = []
        for title, blocks in read:
            if self.numbering and blocks and label_sheets:
                path = outline.enter_level(0)
                title = f"{'.'.join(str(p) for p in path)}. {title}"

            lines: List[str] = []
            for block_lines, headers, is_table in blocks:
                if self.numbering and is_table:
                    block_lines = number_table_lines(
                        block_lines, outline.next_table(), headers
                    )
                lines.extend(block_lines)
                total_tables += int(is_table)
            sheets.append((title, lines))
            all_lines.extend(lines)

        write_docx(sheets, output_path, label_sheets)

        summary: Dict[str, Any] = {
            "input_file": input_path,
            "output_file": output_path,
            "sheets_read_count": len(sheets),
            "pages_patched_count": 0,
            "total_tables_flattened": total_tables,
            "continuation_pages_added": 0,
            "status": "success",
        }

        if has_uncached_formulas(input_path):
            summary["uncached_formulas"] = True
            logger.warning(
                "%s holds formulas Excel has never calculated; those cells are "
                "empty in the workbook itself and so are empty in the output. "
                "Open the file in Excel and save it to store the results.",
                input_path,
            )

        if self.verify_output:
            from .verifier import verify_xlsx

            report = verify_xlsx(input_path, output_path, all_lines)
            summary["verification_passed"] = report.passed
            summary["verification"] = report
            if not report.passed:
                summary["status"] = "verification_failed"
                logger.warning("Verification failed:\n%s", report.describe())
            else:
                logger.info("Verification passed:\n%s", report.describe())

        logger.info(
            "Done: %s", {k: v for k, v in summary.items() if k != "verification"}
        )
        return summary
