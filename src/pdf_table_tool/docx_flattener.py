"""Flatten every table of a Word document into bullet paragraphs, in place.

The PDF path has to reconstruct a table from ink on a page; a .docx already
states its own structure, so this module reads it directly -- cell spans from
``w:gridSpan``/``w:vMerge``, list levels from ``w:numPr`` -- and hands the same
:class:`Grid` to the same :class:`TableFormatter`.  Both formats therefore
produce identical bullets for identical tables.

The table element is then replaced, at its exact position in the document, by
one paragraph per bullet line, so everything around it (headings, images,
section layout, headers/footers) is untouched.
"""

import logging
import re
from dataclasses import dataclass
from typing import Any, Dict, Iterator, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from .config import settings
from .formatter import (
    INDENT_UNIT,
    MAX_LABEL_CHARS,
    TableFormatter,
    TableStructure,
    _has_label_column,
    _logical_rows,
    detect_structure,
    normalise_sliced_cells,
)
from .docx_numbering import ListNumbering, numbering_of
from .grid_extractor import CellLine, Grid, GridCell, Item
from .outline import (
    DocumentOutline,
    keep_unsaid_headers,
    number_table_lines,
    parse_heading,
    table_title,
)
from .text_utils import (
    bullet_marker,
    is_bullet_line,
    join_wrapped_lines,
    normalize_text,
    strip_bullet,
)

logger = logging.getLogger(__name__)

# Synthetic geometry.  Nothing here is measured from the document: a .docx cell
# carries its items explicitly, so these numbers only have to keep the
# geometric helpers (cell slicing, column edges) from firing.  Rows are spaced
# with a gap wider than the 3pt "touching" tolerance so `normalise_sliced_cells`
# never fuses two genuine Word rows.
COL_WIDTH = 120.0
ROW_HEIGHT = 40.0
ROW_GAP = 12.0
LINE_HEIGHT = 10.0

# How far one bullet level is indented in the rendered Word paragraph.
INDENT_STEP_PT = 18.0

MAX_LEVEL = 3


# ---------------------------------------------------------------------------
# reading the table
# ---------------------------------------------------------------------------
def _int_attr(element, attr: str, default: int = 0) -> int:
    if element is None:
        return default
    raw = element.get(qn(attr))
    try:
        return int(raw)
    except (TypeError, ValueError):
        return default


def _grid_span(tc) -> int:
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return 1
    return max(1, _int_attr(tc_pr.find(qn("w:gridSpan")), "w:val", 1))


def _vmerge(tc) -> Optional[str]:
    """``"restart"``, ``"continue"`` or None -- the cell's vertical merge role."""
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    merge = tc_pr.find(qn("w:vMerge"))
    if merge is None:
        return None
    return merge.get(qn("w:val")) or "continue"


def _paragraph_blocks(p) -> List[str]:
    """A paragraph's visible text, split at manual line breaks.

    Text is collected from every ``w:t`` descendant rather than from
    ``Paragraph.text`` so runs inside hyperlinks, fields and content controls
    are not silently dropped.
    """
    parts: List[str] = []
    for node in p.iter():
        tag = node.tag
        if tag == qn("w:t"):
            parts.append(node.text or "")
        elif tag == qn("w:tab"):
            parts.append(" ")
        elif tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
        elif tag == qn("w:noBreakHyphen"):
            parts.append("-")
    blocks = [normalize_text(b) for b in "".join(parts).split("\n")]
    return [b for b in blocks if b]


def _outline_level(p_pr) -> Optional[int]:
    """The heading rank one ``w:pPr`` block states, if any.

    ``w:val="9"`` is Word's way of saying "body text", so only 0-8 count.
    """
    if p_pr is None:
        return None
    level = p_pr.find(qn("w:outlineLvl"))
    if level is None:
        return None
    value = _int_attr(level, "w:val", 9)
    return value if 0 <= value <= 8 else None


def _pr_level(p_pr) -> Tuple[Optional[int], bool]:
    """``(level, is_list)`` contributed by one ``w:pPr`` block, if any."""
    if p_pr is None:
        return None, False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        return min(MAX_LEVEL, _int_attr(num_pr.find(qn("w:ilvl")), "w:val", 0)), True
    ind = p_pr.find(qn("w:ind"))
    if ind is not None:
        left = _int_attr(ind, "w:left", 0) or _int_attr(ind, "w:start", 0)
        # 360 twips = 0.25", the step Word uses for one indent click.
        return min(MAX_LEVEL, max(0, left // 360)), False
    return None, False


# Word's built-in heading styles keep the same ids in every language -- the
# *name* is translated, the id is not -- so the id is what identifies them.
_HEADING_STYLE_RE = re.compile(r"^Heading\s*([1-9])$", re.IGNORECASE)


class StyleIndex:
    """What each paragraph style says about list level, indent and outline rank.

    "List Bullet 2" carries its numbering in styles.xml, not on the paragraph,
    so reading only ``w:pPr`` would flatten every styled list to one level.  The
    same is true of headings: "Heading 2" is what makes a paragraph a section
    title, and the paragraph itself says nothing about it.
    """

    def __init__(self, styles_element=None):
        self._levels: Dict[str, Tuple[Optional[int], bool]] = {}
        self._outline: Dict[str, int] = {}
        self._lists: Dict[str, Tuple[str, int]] = {}
        self._based_on: Dict[str, str] = {}
        if styles_element is None:
            return
        for style in styles_element.findall(qn("w:style")):
            style_id = style.get(qn("w:styleId"))
            if not style_id:
                continue
            p_pr = style.find(qn("w:pPr"))
            self._levels[style_id] = _pr_level(p_pr)
            outline = _outline_level(p_pr)
            if outline is None:
                heading = _HEADING_STYLE_RE.match(style_id)
                outline = int(heading.group(1)) - 1 if heading else None
            if outline is not None:
                self._outline[style_id] = outline
            reference = ListNumbering.reference(p_pr) if p_pr is not None else None
            if reference is not None and reference[0] is not None:
                self._lists[style_id] = (reference[0], reference[1] or 0)
            based_on = style.find(qn("w:basedOn"))
            if based_on is not None and based_on.get(qn("w:val")):
                self._based_on[style_id] = based_on.get(qn("w:val"))

    def level_of(self, style_id: Optional[str]) -> Tuple[Optional[int], bool]:
        seen: set = set()
        while style_id and style_id in self._levels and style_id not in seen:
            seen.add(style_id)
            level, is_list = self._levels[style_id]
            if level is not None:
                return level, is_list
            style_id = self._based_on.get(style_id)
        return None, False

    def outline_of(self, style_id: Optional[str]) -> Optional[int]:
        """Heading rank of a style: 0 for "Heading 1", None when it is not one."""
        seen: set = set()
        while style_id and style_id not in seen:
            seen.add(style_id)
            if style_id in self._outline:
                return self._outline[style_id]
            style_id = self._based_on.get(style_id)
        return None

    def list_of(self, style_id: Optional[str]) -> Optional[Tuple[str, int]]:
        """``(numId, ilvl)`` a style carries, as "List Number" and friends do."""
        seen: set = set()
        while style_id and style_id not in seen:
            seen.add(style_id)
            if style_id in self._lists:
                return self._lists[style_id]
            style_id = self._based_on.get(style_id)
        return None


_NO_STYLES = StyleIndex()


def _paragraph_level(p, styles: StyleIndex) -> Tuple[int, bool]:
    """List depth of a paragraph, and whether it is a list item at all.

    Direct formatting wins over the style, as it does in Word itself.
    """
    p_pr = p.find(qn("w:pPr"))
    level, is_list = _pr_level(p_pr)
    if level is not None:
        return level, is_list

    style_id = None
    if p_pr is not None:
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            style_id = p_style.get(qn("w:val"))
    level, is_list = styles.level_of(style_id)
    return (level or 0), is_list


def _style_id(p) -> Optional[str]:
    p_pr = p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    p_style = p_pr.find(qn("w:pStyle"))
    return None if p_style is None else p_style.get(qn("w:val"))


def paragraph_heading(p, styles: StyleIndex = _NO_STYLES):
    """``(number the paragraph states, heading rank of its style)``.

    Both may be present, either, or neither.  A number written into the text
    ("1.1. Đối tượng áp dụng") is taken as it stands, exactly as the PDF path
    reads it -- the same document must number the same way whichever format it
    arrives in.  A heading with no number in its text is one Word draws itself,
    and its rank is all there is to go on.
    """
    blocks = _paragraph_blocks(p)
    if not blocks:
        return None, None
    p_pr = p.find(qn("w:pPr"))
    rank = _outline_level(p_pr)
    if rank is None:
        rank = styles.outline_of(_style_id(p))
    return parse_heading(blocks[0]), rank


def _paragraph_items(p, styles: StyleIndex = _NO_STYLES) -> List[Item]:
    """Turn one Word paragraph into logical bullets."""
    blocks = _paragraph_blocks(p)
    if not blocks:
        return []
    level, numbered = _paragraph_level(p, styles)

    items: List[Item] = []
    for text in blocks:
        if is_bullet_line(text):
            # The author typed the glyph: keep it and drop it from the text so
            # it is not rendered twice.
            items.append(
                Item(level=level, text=strip_bullet(text), marker=bullet_marker(text))
            )
        else:
            items.append(Item(level=level, text=text, marker="-" if numbered else ""))
    return items


def _lines_to_items(lines: List[str]) -> List[Item]:
    """Read finished bullet lines (from a nested table) back as items."""
    items: List[Item] = []
    for line in lines:
        if not line.strip():
            continue
        indent = len(line) - len(line.lstrip(" "))
        level = min(MAX_LEVEL, indent // max(1, len(INDENT_UNIT)))
        text = line.strip()
        marker = bullet_marker(text) if is_bullet_line(text) else ""
        items.append(
            Item(level=level, text=strip_bullet(text) if marker else text, marker=marker)
        )
    return items


def _cell_items(
    tc, formatter: TableFormatter, styles: StyleIndex = _NO_STYLES
) -> List[Item]:
    """Every bullet of a cell, in document order.

    A table drawn inside the cell is flattened first and spliced in at its own
    position, so a sub-table keeps its column pairing instead of dissolving
    into a run of loose words.
    """
    items: List[Item] = []
    for child in tc.iterchildren():
        if child.tag == qn("w:p"):
            items.extend(_paragraph_items(child, styles))
        elif child.tag == qn("w:tbl"):
            try:
                nested_lines = flatten_table_element(child, formatter, styles=styles)
            except Exception as exc:  # pragma: no cover - defensive
                logger.warning("Nested table flattening failed: %s", exc)
                continue
            items.extend(_lines_to_items(nested_lines))
    return items


def _cell_bbox(row: int, col: int, row_span: int, col_span: int):
    return (
        col * COL_WIDTH,
        row * ROW_HEIGHT,
        (col + col_span) * COL_WIDTH - 8.0,
        (row + row_span) * ROW_HEIGHT - ROW_GAP,
    )


def _cell_from_items(
    row: int, col: int, row_span: int, col_span: int, items: List[Item]
) -> GridCell:
    x0, top, x1, bottom = _cell_bbox(row, col, row_span, col_span)
    lines = [
        CellLine(
            text=item.text,
            x0=x0 + 2.0 + item.level * 8.0,
            x1=x0 + 6.0 + item.level * 8.0,
            top=top + 2.0 + i * LINE_HEIGHT,
            bottom=top + 2.0 + i * LINE_HEIGHT + LINE_HEIGHT * 0.8,
        )
        for i, item in enumerate(items)
    ]
    return GridCell(
        row=row,
        col=col,
        row_span=row_span,
        col_span=col_span,
        bbox=(x0, top, x1, bottom),
        lines=lines,
        items=items,
    )


def build_grid_from_table(
    tbl, formatter: TableFormatter, styles: StyleIndex = _NO_STYLES
) -> Grid:
    """Build a :class:`Grid` from a ``w:tbl`` element."""
    rows = tbl.findall(qn("w:tr"))
    if not rows:
        return Grid(n_rows=0, n_cols=0, cells=[])

    cells: List[GridCell] = []
    open_merge: Dict[int, int] = {}  # column -> index of the cell it continues
    n_cols = 1

    for r, tr in enumerate(rows):
        col = 0
        for tc in tr.findall(qn("w:tc")):
            span = _grid_span(tc)
            merge = _vmerge(tc)
            items = _cell_items(tc, formatter, styles)

            if merge == "continue" and col in open_merge:
                anchor = cells[open_merge[col]]
                anchor.row_span = r - anchor.row + 1
                x0, top, x1, _bottom = anchor.bbox
                anchor.bbox = (
                    x0,
                    top,
                    x1,
                    (anchor.row + anchor.row_span) * ROW_HEIGHT - ROW_GAP,
                )
                # A continuation cell is normally empty, but Word does let text
                # live there -- it belongs to the merged cell.
                if items:
                    anchor.items = (anchor.items or []) + items
                    anchor.lines.extend(
                        _cell_from_items(anchor.row, anchor.col, 1, span, items).lines
                    )
            else:
                cells.append(_cell_from_items(r, col, 1, span, items))
                if merge == "restart":
                    open_merge[col] = len(cells) - 1
                else:
                    open_merge.pop(col, None)
            col += span
        n_cols = max(n_cols, col)

    cells.sort(key=lambda c: (c.row, c.col))
    return Grid(
        n_rows=len(rows),
        n_cols=n_cols,
        cells=cells,
        bbox=(0.0, 0.0, n_cols * COL_WIDTH, len(rows) * ROW_HEIGHT),
    )


# ---------------------------------------------------------------------------
# header band -- the one thing a real document knows and a PDF does not
# ---------------------------------------------------------------------------
# The shared geometry rule refuses to read row 0 as a header when it merges
# columns the rows below keep apart, because on a page that pattern is usually
# a labelled data row and there is no way to tell.  A .docx does not have that
# ambiguity: the author merged the cell on purpose, and Word even records which
# rows repeat as the header of a long table.  So an editable format gets its own
# pass, run *before* the shared formatter.  Excel states both facts too (merged
# ranges, print titles), so :mod:`.xlsx_flattener` uses this pass as well --
# nothing below reads a Word element, only a :class:`Grid` and a row count.


def _declared_header_rows(tbl) -> int:
    """Leading rows the author ticked as "Repeat as header row"."""
    band = 0
    for tr in tbl.findall(qn("w:tr")):
        tr_pr = tr.find(qn("w:trPr"))
        header = None if tr_pr is None else tr_pr.find(qn("w:tblHeader"))
        if header is None or header.get(qn("w:val")) in ("0", "false"):
            break
        band += 1
    return band


def _grid_row(grid: Grid, row: int) -> List[GridCell]:
    return sorted((c for c in grid.cells if c.row == row), key=lambda c: c.col)


def _reads_as_labels(cells: List[GridCell]) -> bool:
    """Could this row be naming columns rather than holding their values?"""
    filled = [c for c in cells if not c.is_empty]
    if not filled:
        return False
    for cell in filled:
        text = " ".join(cell.text.split())
        # A sentence is a value, not a column name -- this is what keeps a
        # two-dimensional table (labels down the side, a paragraph merged
        # across the top) from being read upside down.
        if len(text) > MAX_LABEL_CHARS or is_bullet_line(text):
            return False
    if all(re.fullmatch(r"[\d.,%/\s-]+", c.text.strip()) for c in filled):
        return False
    return True


def declared_header_band(grid: Grid, declared: int) -> int:
    """How many leading rows label the columns, using what the source states.

    `declared` is the header row count the author set outright -- Word's "Repeat
    as header row", Excel's "Rows to repeat at top".  Returns 0 when the source
    adds nothing: the shared geometric rule then decides on its own, exactly as
    it does for a PDF.
    """
    if grid.n_rows < 2 or grid.n_cols < 2:
        return 0

    row0 = _grid_row(grid, 0)
    if not _reads_as_labels(row0):
        return 0

    band = min(declared, grid.n_rows - 1, 3)
    if band >= 1:
        # Declared outright -- a column the author left unnamed is their choice,
        # not evidence against the header.
        return band

    # Nothing was declared, so the merge itself has to carry the case: row 0
    # names every column, and does it by spanning some of them.  A row naming
    # only part of the table is a data row, not a header.
    covered = set()
    for cell in row0:
        if not cell.is_empty:
            covered.update(range(cell.col, cell.col + cell.col_span))

    merged = [c for c in row0 if c.col_span > 1 and not c.is_empty]
    if not merged or len(covered) < grid.n_cols:
        return 0

    # A merged cell over sub-labels is a two-row header ("Tỷ lệ" over "Cơ bản"
    # and "Đóng thêm").  It only counts as one when the second row is made of
    # labels too and sits entirely under the merged cells.
    if grid.n_rows >= 3:
        row1 = [c for c in _grid_row(grid, 1) if not c.is_empty]
        merged_cols = {
            c for cell in merged for c in range(cell.col, cell.col + cell.col_span)
        }
        if (
            len(row1) >= 2
            and _reads_as_labels(row1)
            and all(c.col in merged_cols for c in row1)
        ):
            return 2
    return 1


def fuse_header_band(grid: Grid, band: int) -> Grid:
    """Rewrite the header band as one plain row of one label per column.

    Handing the shared formatter a header row without merges is what lets it
    do its job: it pairs each value with the label above it, and the merged
    cell's text reaches every column it covered instead of only the first.
    """
    labels = ["" for _ in range(grid.n_cols)]
    for row in range(band):
        for cell in _grid_row(grid, row):
            if cell.is_empty:
                continue
            text = join_wrapped_lines([ln.text for ln in cell.lines])
            for col in range(cell.col, min(cell.col + cell.col_span, grid.n_cols)):
                labels[col] = f"{labels[col]} {text}".strip() if labels[col] else text

    cells = [
        _cell_from_items(0, col, 1, 1, [Item(level=0, text=text)] if text else [])
        for col, text in enumerate(labels)
    ]

    shift = band - 1
    for cell in grid.cells:
        if cell.row < band:
            continue
        row = cell.row - shift
        cells.append(
            GridCell(
                row=row,
                col=cell.col,
                row_span=cell.row_span,
                col_span=cell.col_span,
                bbox=_cell_bbox(row, cell.col, cell.row_span, cell.col_span),
                lines=cell.lines,
                items=cell.items,
            )
        )

    cells.sort(key=lambda c: (c.row, c.col))
    n_rows = grid.n_rows - shift
    return Grid(
        n_rows=n_rows,
        n_cols=grid.n_cols,
        cells=cells,
        bbox=(0.0, 0.0, grid.n_cols * COL_WIDTH, n_rows * ROW_HEIGHT),
    )


def declared_header_structure(
    grid: Grid, declared: int
) -> Tuple[Grid, Optional[TableStructure]]:
    """Apply what the source states about the header band.

    Returns the grid to format and, when the source had something to add, the
    layout to format it with.  A None structure means it knew nothing the shared
    geometric rule does not already handle, so the PDF path's own analysis is
    used unchanged.
    """
    band = declared_header_band(grid, declared)
    if band == 0 or band <= detect_structure(grid).header_rows:
        # Geometry already found it (or there is nothing to find); leave a
        # working table alone.
        return grid, None

    logger.debug("Word header band of %d row(s) applied", band)
    grid = fuse_header_band(grid, band)

    structure = detect_structure(grid)
    if structure.header_rows >= 1:
        # With the merges resolved the shared rule agrees on its own.
        return grid, structure

    # It still refuses -- for a reason that only holds on a page, such as the
    # header naming fewer columns than the rows below fill.  Word said this row
    # is the header, so say so.
    matrix_cells = {(c.row, c.col): c for c in grid.cells}
    rows = _logical_rows(grid, matrix_cells, 1)
    return grid, TableStructure(
        header_rows=1,
        label_column=_has_label_column(rows, grid.n_cols),
        source="docx",
    )


def flatten_table_element(
    tbl,
    formatter: TableFormatter,
    styles: StyleIndex = _NO_STYLES,
    outline: Optional[DocumentOutline] = None,
    section: Optional[Tuple[int, ...]] = None,
    title: str = "",
    alone: bool = False,
) -> List[str]:
    """Bullet lines for one ``w:tbl`` element (nested tables included).

    With an `outline`, the table is given its place in the document's numbering
    and its rows are numbered under it; `section` is the branch it belongs to,
    `title` the name the paragraph above gave it and `alone` whether the section
    holds nothing else, all worked out beforehand by :func:`plan_sections`.  A
    table drawn inside a cell is called without any of them: it is part of its
    parent's row, not a section of its own.
    """
    grid = build_grid_from_table(tbl, formatter, styles)
    if grid.n_rows == 0:
        return []

    grid, structure = declared_header_structure(grid, _declared_header_rows(tbl))
    normalised = normalise_sliced_cells(grid)
    if structure is None:
        structure = detect_structure(normalised)

    lines, headers = formatter.format_grid(grid, None, structure)
    if outline is not None and lines:
        # The number is taken only now: a table that turned out to be empty
        # would otherwise leave a hole in the numbering.
        number = outline.next_table(section, alone)
        # A table that took the section's own number is already named by the
        # heading standing above it, under that very number -- captioning it
        # here would print that line a second time.
        named = "" if number.path == tuple(section or ()) else title
        lines = number_table_lines(lines, number, named, headers)
    elif lines:
        # No caption line to carry the headers no row repeats, so they get a
        # line of their own -- or they are lost.
        lines = keep_unsaid_headers(lines, headers)
    return lines


def _follow_heading(
    outline: DocumentOutline,
    p,
    styles: StyleIndex,
    listed: Optional[Tuple[int, ...]] = None,
) -> None:
    """Move the outline on if this paragraph is a heading.

    Three sources, in the order of how much they know:

    1. a number written into the text ("1.1. Đối tượng áp dụng") -- the author
       typed it, and it is what the PDF path would read off the page;
    2. the number Word draws from its own list counters, which is what the
       reader sees for a document that numbers itself;
    3. the rank of a heading style, counted here, for a document that uses
       Heading 1/2/3 and lets Word draw nothing at all.
    """
    path, rank = paragraph_heading(p, styles)
    if path is not None:
        outline.enter(path)
    elif listed is not None:
        outline.enter(listed)
    elif rank is not None:
        outline.enter_level(rank)


def _paragraph_title(p, styles: StyleIndex, listed=None) -> Optional[str]:
    """The name this paragraph gives the table below it, or None if it is blank.

    None and "" are different answers: a blank paragraph between a caption and
    its table says nothing, so the caption above it still stands, whereas a
    paragraph of prose has had its say and the table is left unnamed.

    The last block is the one asked, because a paragraph broken by a manual line
    break ends on the line that actually sits above the table.
    """
    blocks = _paragraph_blocks(p)
    if not blocks:
        return None
    path, rank = paragraph_heading(p, styles)
    # A heading is a heading whichever way this document writes one down: in the
    # text, on a Heading style, or through a list counter Word draws itself.
    heading = path is not None or rank is not None or listed is not None
    return table_title(blocks[-1], heading=heading)


@dataclass
class TablePlan:
    """Where one table sits, worked out before the body is touched."""

    # Kept so lxml cannot recycle the element's id onto an unrelated object.
    element: Any
    section: Tuple[int, ...]
    title: str
    # The section holds this table and no other, so the table is the section.
    alone: bool = False


def plan_sections(
    document, styles: StyleIndex, outline: DocumentOutline
) -> Dict[int, TablePlan]:
    """Walk the body once and work out which section each table belongs to.

    Word's list counters are the reason this is a pass of its own: they only
    come out right when *every* numbered paragraph is counted in reading order,
    including the ones inside tables, and by the time a table is being replaced
    its own paragraphs are already gone.  The same walk sees the paragraph above
    each table, which is what names it, and counts how many tables each section
    holds -- a section holding just the one is that table (see
    :meth:`DocumentOutline.next_table`), and that cannot be known until the last
    table of the body has been seen.

    Returns ``{id(table element): TablePlan}``.
    """
    numbering = numbering_of(document)
    plan: Dict[int, TablePlan] = {}
    tables_in: Dict[Tuple[int, ...], int] = {}
    title = ""
    for node in _walk(document):
        if node.tag == qn("w:p"):
            listed = numbering.advance(node, styles.list_of(_style_id(node)))
            if not _inside_table(node):
                _follow_heading(outline, node, styles, listed)
                said = _paragraph_title(node, styles, listed)
                if said is not None:
                    title = said
        elif not _inside_table(node):
            plan[id(node)] = TablePlan(node, outline.section, title)
            tables_in[outline.section] = tables_in.get(outline.section, 0) + 1
            # One table does not caption the next one.
            title = ""

    for planned in plan.values():
        planned.alone = tables_in[planned.section] == 1
    return plan


# ---------------------------------------------------------------------------
# writing the bullets back
# ---------------------------------------------------------------------------
def _table_typography(tbl) -> Tuple[Optional[str], Optional[float]]:
    """Font face and size of the first real run, so the bullets match the table."""
    name: Optional[str] = None
    size: Optional[float] = None
    for run in tbl.iter(qn("w:r")):
        if not any((t.text or "").strip() for t in run.findall(qn("w:t"))):
            continue
        r_pr = run.find(qn("w:rPr"))
        if r_pr is not None:
            fonts = r_pr.find(qn("w:rFonts"))
            if fonts is not None:
                name = fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi"))
            half_points = _int_attr(r_pr.find(qn("w:sz")), "w:val", 0)
            if half_points:
                size = max(settings.MIN_FONT_SIZE, min(14.0, half_points / 2.0))
        if name or size:
            break
    return name, size


def _apply_font(run, name: Optional[str], size: Optional[float]) -> None:
    if size:
        run.font.size = Pt(size)
    if not name:
        return
    run.font.name = name
    r_pr = run._r.get_or_add_rPr()
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:  # pragma: no cover - get_or_add via font.name always adds
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    # Vietnamese diacritics are laid out with the complex-script face too.
    fonts.set(qn("w:cs"), name)


def replace_table_with_bullets(tbl, lines: List[str], container) -> None:
    """Swap a ``w:tbl`` for one paragraph per bullet line, at the same spot."""
    name, size = _table_typography(tbl)
    added = 0
    for line in lines:
        if not line.strip():
            # The blank line between two records earns its paragraph: it is what
            # a chunker reads as the end of one record and the start of the
            # next, and what keeps the list from running together on the page.
            if added:
                tbl.addprevious(OxmlElement("w:p"))
                added += 1
            continue
        indent = len(line) - len(line.lstrip(" "))
        level = indent // max(1, len(INDENT_UNIT))
        p = OxmlElement("w:p")
        tbl.addprevious(p)
        paragraph = Paragraph(p, container)
        fmt = paragraph.paragraph_format
        if level:
            fmt.left_indent = Pt(INDENT_STEP_PT * level)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(2)
        _apply_font(paragraph.add_run(line.strip()), name, size)
        added += 1

    parent = tbl.getparent()
    if added == 0 and not parent.findall(qn("w:p")):
        # A cell (and the document body) must not end up without a paragraph.
        tbl.addprevious(OxmlElement("w:p"))
    parent.remove(tbl)


# ---------------------------------------------------------------------------
# pipeline
# ---------------------------------------------------------------------------
def _containers(document) -> Iterator[Any]:
    """Every story that can hold a table: the body, then headers and footers."""
    yield document
    for section in document.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if part is not None:
                yield part


def _walk(container) -> Iterator[Any]:
    """Every paragraph and table of a story, in reading order, nesting included.

    ``container.tables`` only reports tables sitting directly in the story, so a
    table inside a text box or a content control would survive untouched; this
    walks the whole tree instead.
    """
    element = getattr(container, "_element", None)
    if element is None:
        element = container.element
    return element.iter(qn("w:p"), qn("w:tbl"))


def _inside_table(node) -> bool:
    parent = node.getparent()
    while parent is not None:
        if parent.tag == qn("w:tbl"):
            return True
        parent = parent.getparent()
    return False


def _blocks_in_order(container) -> Iterator[Any]:
    """Paragraphs and outermost tables of a story, in reading order.

    A table's own paragraphs are its content, and a table drawn inside a cell is
    reached through its parent -- so both are left out here.

    Reading order is what makes the headings usable: a table has to be numbered
    under the heading that precedes it, so the two must arrive interleaved.
    """
    return (node for node in _walk(container) if not _inside_table(node))


class DocxTableFlattener:
    """Same contract as :class:`PDFTableFlattenerPipeline`, for .docx files."""

    def __init__(self, verify_output: bool = True, numbering: bool = True):
        self.formatter = TableFormatter()
        self.verify_output = verify_output
        self.numbering = numbering

    def process(self, input_path: str, output_path: str) -> Dict[str, Any]:
        logger.info("Processing %s", input_path)
        document = Document(input_path)
        styles = StyleIndex(document.styles.element)

        outline = DocumentOutline()
        # One pass over the whole body before anything is touched: it settles
        # which section each table belongs to, and it reserves every number the
        # document itself uses, so a table is never handed one that a heading
        # further down owns.
        plan = plan_sections(document, styles, outline) if self.numbering else {}

        total_tables = 0
        all_lines: List[str] = []
        # Keyed by id, but holding the element itself: lxml hands out its proxy
        # objects on demand and recycles their ids once they are collected, so a
        # bare set of ids would drop unrelated tables as "already seen".
        seen: Dict[int, Any] = {}

        for container in _containers(document):
            # Word reuses one header object across linked sections; the table
            # elements are then literally the same, so flatten them once.
            # The list is materialised before anything is replaced: removing a
            # table from a live tree derails an iterator still walking it, and
            # every table after the first would be dropped.
            for node in [n for n in _blocks_in_order(container) if n.tag == qn("w:tbl")]:
                if id(node) in seen:
                    continue
                seen[id(node)] = node
                # A table in a page header has no place in the document's
                # outline, and neither has one the planning pass never saw:
                # both are numbered as sections of their own.
                planned = plan.get(id(node)) or TablePlan(node, (), "")
                lines = flatten_table_element(
                    node,
                    self.formatter,
                    styles,
                    outline if self.numbering else None,
                    planned.section,
                    planned.title,
                    planned.alone,
                )
                replace_table_with_bullets(node, lines, container)
                if lines:
                    total_tables += 1
                    all_lines.extend(lines)

        document.save(output_path)

        summary: Dict[str, Any] = {
            "input_file": input_path,
            "output_file": output_path,
            "pages_patched_count": 0,
            "total_tables_flattened": total_tables,
            "continuation_pages_added": 0,
            "status": "success",
        }

        if self.verify_output:
            from .verifier import verify_docx

            report = verify_docx(input_path, output_path, all_lines)
            summary["verification_passed"] = report.passed
            summary["verification"] = report
            if not report.passed:
                summary["status"] = "verification_failed"
                logger.warning("Verification failed:\n%s", report.describe())
            else:
                logger.info("Verification passed:\n%s", report.describe())

        logger.info(
            "Done: %s", {k: v for k, v in summary.items() if k != "verification"}
        )
        return summary
