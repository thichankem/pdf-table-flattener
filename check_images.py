import os
import sys
import docx
from pdf2docx import Converter

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

pdf_files = [
    "Bảo hiểm nhân thọ - Lộc Phát Hưng Thịnh - Quy định sản phẩm.pdf",
    "Bảo hiểm nhân thọ - Lộc Phát Tràng An - Quy định sản phẩm.pdf",
    "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf",
    "Sản phẩm cho vay kinh doanh xi măng Xuân Thành - Kênh quầy.docx.pdf",
    "Sản phẩm cho vay tái tài trợ - kênh quầy.docx.pdf"
]

for idx, pdf_name in enumerate(pdf_files, 1):
    print(f"\n==================================================")
    print(f"CHECKING IMAGES FOR FILE {idx}: {pdf_name}")
    print(f"==================================================")
    
    temp_docx = f"temp_orig_{idx}.docx"
    cv = Converter(pdf_name)
    cv.convert(temp_docx)
    cv.close()
    
    doc = docx.Document(temp_docx)
    
    # Count inline shapes in doc
    inline_shapes_count = len(doc.inline_shapes)
    
    # Count w:drawing and w:pict XML tags in document body
    xml_str = doc._body._element.xml
    drawing_count = xml_str.count('w:drawing')
    pict_count = xml_str.count('w:pict')
    
    print(f"Initial DOCX extracted by pdf2docx:")
    print(f"  • Inline Shapes count: {inline_shapes_count}")
    print(f"  • XML <w:drawing> count: {drawing_count}")
    print(f"  • XML <w:pict> count: {pict_count}")
    
    # Check if images are inside tables
    table_images = 0
    for t_idx, table in enumerate(doc.tables, 1):
        tbl_xml = table._element.xml
        if 'w:drawing' in tbl_xml or 'w:pict' in tbl_xml:
            table_images += 1
            print(f"  ⚠️ Table {t_idx} contains image/drawing! (Rows: {len(table.rows)}, Cols: {len(table.columns)})")
            
    print(f"  • Tables containing images: {table_images} / {len(doc.tables)}")
    
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
