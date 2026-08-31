"""Hierarchical numbering: a table takes its place in the document's outline.

The point of it is chunking for retrieval, so the tests are about the two things
a chunker needs: a line must say which branch it belongs to, and that branch must
be the one the document itself would give it -- a table under "1. Thuật ngữ" is
1.1 and its rows 1.1.1, never a numbering of its own invention.
"""

import fitz
import pytest
from docx import Document
from docx.oxml import parse_xml
from openpyxl import Workbook

from pdf_table_tool.config import settings
from pdf_table_tool.docx_flattener import DocxTableFlattener
from pdf_table_tool.docx_numbering import ListNumbering
from pdf_table_tool.outline import (
    CONTINUED_MARK,
    DocumentOutline,
    TableNumber,
    continues_outline,
    is_contents_line,
    number_table_lines,
    parse_heading,
    row_context,
    table_title,
)
from pdf_table_tool.pipeline import PDFTableFlattenerPipeline
from pdf_table_tool.xlsx_flattener import XlsxTableFlattener


# ------------------------------------------------------------ reading headings
@pytest.mark.parametrize(
    "line, expected",
    [
        ("1. THUẬT NGỮ VÀ ĐỊNH NGHĨA", (1,)),
        ("1.1. Bác sĩ: là người có bằng cấp chuyên môn", (1, 1)),
        ("2.3.2.Điều khoản cung cấp thông tin", (2, 3, 2)),  # no blank after the dot
        ("1.27 Thương tật toàn bộ và vĩnh viễn", (1, 27)),
        ("3) Quyền lợi bảo hiểm", (3,)),
        ("ĐIỀU 5: QUYỀN LỢI BẢO HIỂM", (5,)),
        ("ĐIỀU5:QUYỀN LỢI", (5,)),  # letter-spaced heading, blanks lost
    ],
)
def test_a_numbered_heading_is_read(line, expected):
    assert parse_heading(line) == expected


@pytest.mark.parametrize(
    "line",
    [
        "1.000.000 đồng phí bảo hiểm",  # a sum of money, not section 1.000.000
        "3 Bản sao là bản sao y chứng thực",  # a footnote: no punctuation
        "3.5% của Số tiền bảo hiểm",
        "theo quy định tại Điều 5 của Hợp đồng",  # a cross-reference, mid-sentence
        "Điều 8.2 quy định về tạm ứng",  # a reference to a clause, not article 8
        "4/23",  # a page number
        "",
    ],
)
def test_prose_is_not_read_as_a_heading(line):
    assert parse_heading(line) is None


def test_a_contents_entry_is_recognised_by_its_leader_dots():
    assert is_contents_line("2.1. Quyền lợi ....................... 7")
    assert not is_contents_line("2.1. Quyền lợi có thể được bảo hiểm")


# ------------------------------------------------------- the name of a table
@pytest.mark.parametrize(
    "line, heading, expected",
    [
        # A lead-in: the sentence hands over to the table below it.
        ("Chức danh lãnh đạo và phân hạng tương ứng như bảng sau:",
         False, "Chức danh lãnh đạo và phân hạng tương ứng như bảng sau"),
        # The document says outright that this is a table.
        ("Bảng 3: Biểu phí bảo hiểm", False, "Bảng 3: Biểu phí bảo hiểm"),
        # A heading, whichever way the document numbers it.
        ("2.3. Nguyên tắc xác định đơn vị quản lý",
         True, "Nguyên tắc xác định đơn vị quản lý"),
        ("ĐIỀU 5: QUYỀN LỢI BẢO HIỂM", True, "QUYỀN LỢI BẢO HIỂM"),
        ("Phân hạng và tiêu chí xếp hạng", True, "Phân hạng và tiêu chí xếp hạng"),
    ],
)
def test_the_line_above_a_table_names_it(line, heading, expected):
    assert table_title(line, heading=heading) == expected


@pytest.mark.parametrize(
    "line, heading",
    [
        # Prose is about the table; it is not its name.
        ("Việc phân bổ đơn vị quản lý được thực hiện theo nguyên tắc dưới đây.", False),
        # A styled paragraph that is still a whole sentence.
        ("Khách hàng đáp ứng điều kiện dưới đây được định danh là KHCNUT.", True),
        ("2.1. Quyền lợi ....................... 7", True),  # a contents entry
        ("Đơn vị: triệu đồng " + "x" * 200 + ":", True),  # a paragraph, not a name
        ("", True),
        ("2.3.", True),  # a number and nothing else
    ],
)
def test_a_line_that_names_no_table_gives_no_caption(line, heading):
    assert table_title(line, heading=heading) == ""


def test_an_outline_carries_on_by_one_step():
    assert continues_outline((1, 12), (1, 13))
    assert continues_outline((1, 28), (2,))
    assert continues_outline((2,), (2, 1))
    assert not continues_outline((1, 28), (2, 1))
    assert not continues_outline((1, 12), (1, 20))


# --------------------------------------------------------- allocating a number
def test_a_table_is_numbered_under_the_heading_above_it():
    outline = DocumentOutline()
    outline.enter((1,))
    assert outline.next_table().label == "1.1"
    assert outline.next_table().label == "1.2"


def test_a_table_never_takes_a_number_a_heading_owns():
    """The document has its own 1.1 and 1.2 further down; the table gets 1.3."""
    outline = DocumentOutline()
    outline.load([], known=[(1,), (1, 1), (1, 2)])
    outline.enter((1,))
    assert outline.next_table().label == "1.3"


def test_a_document_without_headings_numbers_its_tables_in_order():
    outline = DocumentOutline()
    assert outline.next_table().label == "1"
    assert outline.next_table().label == "2"


def test_a_deeper_heading_takes_the_table_deeper_with_it():
    outline = DocumentOutline()
    outline.enter((2,))
    outline.enter((2, 3))
    outline.enter((2, 3, 1))
    assert outline.next_table().label == "2.3.1.1"


def test_a_section_holding_one_table_is_that_table():
    """3.5.1 under a heading that already says 3.5 spends a level saying nothing."""
    outline = DocumentOutline()
    outline.enter((3, 5))
    assert outline.next_table(alone=True).label == "3.5"


def test_a_section_that_numbers_children_of_its_own_still_opens_a_level():
    """Rows numbered 3.5.1 would take a number a real heading further down owns."""
    outline = DocumentOutline()
    outline.load([], known=[(3, 5), (3, 5, 1)])
    outline.enter((3, 5))
    assert outline.next_table(alone=True).label == "3.5.2"


def test_a_table_outside_every_section_has_no_number_to_keep():
    outline = DocumentOutline()
    assert outline.next_table(alone=True).label == "1"


def test_a_heading_with_no_number_of_its_own_is_counted():
    """Word draws the number of a Heading style, so it is nowhere in the text."""
    outline = DocumentOutline()
    assert outline.enter_level(0) == (1,)
    assert outline.enter_level(1) == (1, 1)
    assert outline.enter_level(1) == (1, 2)
    assert outline.enter_level(0) == (2,)


# ------------------------------------------------------------ numbering lines
def test_rows_are_numbered_one_level_and_nothing_deeper_is():
    """The numbering stops at the row; what nests in it keeps its own marker.

    A short row is never cut away from the lines under it, so those lines say
    nothing twice -- the row above them is in the same chunk.
    """
    lines = [
        "- Thuật ngữ: Bên mua bảo hiểm  |  Định nghĩa: Tổ chức, cá nhân",
        "- Thuật ngữ: Người được bảo hiểm  |  Định nghĩa:",
        "  - Là người được nhận quyền lợi",
        "  - Phải cư trú tại Việt Nam",
    ]
    out = number_table_lines(lines, TableNumber(path=(1, 1)), "Giải thích từ ngữ")

    assert out == [
        "1.1 Giải thích từ ngữ",
        "",
        "  1.1.1 Thuật ngữ: Bên mua bảo hiểm  |  Định nghĩa: Tổ chức, cá nhân",
        "",
        "  1.1.2 Thuật ngữ: Người được bảo hiểm  |  Định nghĩa:",
        "    - Là người được nhận quyền lợi",
        "    - Phải cư trú tại Việt Nam",
    ]


def test_records_are_separated_by_a_blank_line():
    """A chunker cuts at a blank line before it cuts through a sentence."""
    out = number_table_lines(["- Nam  |  25", "- Lan  |  30"], TableNumber(path=(1,)))
    assert out == ["1.1 Nam  |  25", "", "1.2 Lan  |  30"]


def test_a_table_the_document_never_named_gets_no_caption_line():
    """Column headers are not a name: every row already says them."""
    out = number_table_lines(
        ["- Năm: 1  |  Tỷ lệ: 50%"],
        TableNumber(path=(3,)),
        headers=["Năm", "Tỷ lệ"],
    )
    assert out == ["3.1 Năm: 1  |  Tỷ lệ: 50%"]


def test_a_header_no_row_says_is_kept_so_its_words_are_not_lost():
    """A header the rows never repeat is written nowhere else in the document."""
    out = number_table_lines(
        ["- Ưu tiên 1: Tổng số dư cao nhất"],
        TableNumber(path=(3,)),
        headers=["Thứ tự", "Thứ tự"],  # one merged header over two columns
    )
    assert out[0] == "3 Thứ tự"  # named once, not twice
    assert out[-1].strip() == "3.1 Ưu tiên 1: Tổng số dư cao nhất"


def test_a_continued_table_carries_on_where_it_stopped():
    number = TableNumber(path=(2, 1))
    first = number_table_lines(["- a", "- b"], number, "Biểu phí")
    second = number_table_lines(["- c"], number, "Biểu phí", continued=True)

    assert first[-1].strip().startswith("2.1.2 ")
    assert second[0] == f"2.1 Biểu phí {CONTINUED_MARK}"
    assert second[-1].strip().startswith("2.1.3 ")


def test_everything_under_a_row_is_one_flat_list_that_keeps_its_own_markers():
    """Nesting deeper than a row buys nothing a chunker can read.

    The markers the source wrote are what shows the structure instead, so they
    are kept rather than replaced by numbers nobody asked for.
    """
    lines = ["- a", "  - b", "    + c", "      - d"]
    out = number_table_lines(lines, TableNumber(path=(1,)))

    assert out == [
        "1.1 a",
        "  - b",
        "  + c",  # the third level is pulled up to the second
        "  - d",
    ]


def test_a_table_in_a_deep_section_numbers_only_its_caption():
    """The caption names the branch; every line under it is read off the indent."""
    number = TableNumber(path=(2, 3, 1, 1))
    out = number_table_lines(["- Tuổi: 30  |  Phí: 1.000"], number, "Biểu phí")

    assert out == ["2.3.1.1 Biểu phí", "", "  Tuổi: 30  |  Phí: 1.000"]


def test_numbering_only_ever_adds_words_to_a_line():
    lines = ["- Khoản: 3.8  |  Quy định: Giải ngân", "  + Chuyển khoản"]
    out = number_table_lines(lines, TableNumber(path=(4,)), "Điều khoản giải ngân")

    body = [ln for ln in out if ln.strip()]
    for source, numbered in zip(lines, body[1:]):
        assert source.strip().lstrip("-+* ") in numbered


@pytest.mark.parametrize(
    "row, expected",
    [
        # The counter is skipped; the label in front of the value is dropped.
        ("1  |  Mục đích: Chuyển tiền học phí  |  Hồ sơ:", "Chuyển tiền học phí"),
        # "Hồ sơ:" announces the column its list fills; it names no record.
        ("Hồ sơ:", ""),
        ("3.8  |  Quy định: Giải ngân", "Giải ngân"),
        # Long enough to be cut, and cut at a word boundary.
        ("Điều kiện: " + "một hai ba bốn năm sáu bảy tám chín mười " * 3,
         "một hai ba bốn năm sáu bảy tám chín mười một hai ba bốn năm"),
    ],
)
def test_the_words_of_a_row_that_name_it(row, expected):
    assert row_context(row) == expected


def test_a_row_that_counts_itself_is_not_counted_twice():
    """"1.3.1 1  |  Mục đích: ..." indexes the row twice and says it once."""
    lines = [
        "- 1  |  Mục đích: Chuyển tiền học phí  |  Hồ sơ:",
        "  - a) Giấy tờ tùy thân của người chuyển tiền",
        "  - b) Hộ chiếu Việt Nam còn hiệu lực",
    ]
    out = number_table_lines(lines, TableNumber(path=(1, 3)))

    assert out == [
        "1.3.1 Mục đích: Chuyển tiền học phí  |  Hồ sơ:",
        "  - a) Giấy tờ tùy thân của người chuyển tiền",
        "  - b) Hộ chiếu Việt Nam còn hiệu lực",
    ]


def test_a_counter_that_disagrees_with_our_count_is_the_document_s_own():
    """Rows 8, 9 of a table continued overleaf keep the numbers they carry."""
    number = TableNumber(path=(1, 3), rows=4)
    out = number_table_lines(["- 8  |  Mục đích: Học phí"], number)

    assert out == ["1.3.5 8  |  Mục đích: Học phí"]


def test_a_deep_row_keeps_its_counter_because_no_number_states_it():
    """Past the numbering limit the row's own counter is the only index it has."""
    out = number_table_lines(
        ["- 1  |  Mục đích: Học phí"], TableNumber(path=(2, 3, 1, 1)), "Hồ sơ"
    )

    assert out == ["2.3.1.1 Hồ sơ", "", "  1  |  Mục đích: Học phí"]


def test_a_divider_between_groups_of_rows_is_not_counted_as_one():
    """A table that counts its rows says, by leaving the count out, "not a row".

    Counting the divider would push every row after it one out of step with the
    count the reader can see -- and that count is what lets the number absorb it.
    """
    lines = [
        "- LOẠI TÀI LIỆU: I. Người cư trú là tổ chức",
        "- 1  |  LOẠI TÀI LIỆU: Thỏa thuận mở tài khoản",
        "- 2  |  LOẠI TÀI LIỆU: Đăng ký mã số thuế",
    ]
    out = number_table_lines(
        lines, TableNumber(path=(2, 1)), headers=["STT", "LOẠI TÀI LIỆU"]
    )

    assert [ln for ln in out if ln.strip()] == [
        "2.1 STT  |  LOẠI TÀI LIỆU",
        "  LOẠI TÀI LIỆU: I. Người cư trú là tổ chức",
        "  2.1.1 LOẠI TÀI LIỆU: Thỏa thuận mở tài khoản",
        "  2.1.2 LOẠI TÀI LIỆU: Đăng ký mã số thuế",
    ]


def test_the_tail_of_a_row_a_page_break_cut_is_not_a_new_row():
    """Numbering the tail would invent a record and misnumber every row after."""
    headers = ["STT", "Nguồn thu"]
    number = TableNumber(path=(1, 4))
    first = number_table_lines(
        ["- 1  |  Nguồn thu: Sổ tiết kiệm", "- 2  |  Nguồn thu: Chuyển nhượng GTCG"],
        number,
        headers=headers,
    )
    second = number_table_lines(
        ["- Nguồn thu: trái phiếu đến hạn", "- 3  |  Nguồn thu: Bán bất động sản"],
        number,
        headers=headers,
        continued=True,
    )

    assert first[-1].strip() == "1.4.2 Nguồn thu: Chuyển nhượng GTCG"
    body = [ln for ln in second if ln.strip()]
    # The tail carries row 2 on, and names it: its own row line is on the page
    # before, so no chunk cut here can reach it.
    assert body[0].strip() == "- Chuyển nhượng GTCG  |  Nguồn thu: trái phiếu đến hạn"
    # The next real row takes the number the document's own count gives it.
    assert body[1].strip() == "1.4.3 Nguồn thu: Bán bất động sản"


def test_a_continuation_page_does_not_print_the_header_row_again():
    """The page the table started on wrote it; eleven copies of it are noise.

    A chunk cut from the eleventh page needs the branch it sits in, and every
    row on it opens with that number itself.
    """
    headers = ["STT", "Mục đích", "Hồ sơ"]
    number = TableNumber(path=(1, 3))
    first = number_table_lines(["- 1  |  Mục đích: Du học"], number, headers=headers)
    second = number_table_lines(
        ["- 2  |  Mục đích: Khám chữa bệnh"], number, headers=headers, continued=True
    )

    # Written once, on the page the table starts on.
    assert first[0] == "1.3 STT  |  Mục đích  |  Hồ sơ"
    assert not any("STT" in line for line in second)
    assert [ln for ln in second if ln.strip()] == ["  1.3.2 Mục đích: Khám chữa bệnh"]


def test_a_long_row_names_its_record_again_before_a_chunk_can_cut_it():
    """A row longer than a chunk is cut in two, and the second half must say so.

    The name comes back every half-chunk, so whichever way the cut falls it is
    inside the same chunk as a copy of it -- rather than on every single line,
    which buys the same thing several times over.
    """
    sentence = ("một hai ba bốn năm sáu bảy tám chín mười " * 4).strip()
    lines = ["- 1  |  Mục đích: Chuyển tiền học phí  |  Hồ sơ:"] + [
        f"  - {sentence}" for _ in range(6)
    ]
    out = number_table_lines(lines, TableNumber(path=(1, 3)))

    nested = [ln for ln in out[1:] if ln.strip()]
    named = [ln for ln in nested if "Chuyển tiền học phí" in ln]
    assert named, "a row long enough to be cut must name itself again"
    assert len(named) * 2 < len(nested), "and not on line after line"


def test_a_row_counters_header_is_written_once_and_never_in_front_of_a_number():
    """"STT: 1" says no more than "1"; the word still belongs to the document."""
    out = number_table_lines(
        ["- 1  |  Mục đích: Học phí", "- 2  |  Mục đích: Sinh hoạt phí"],
        TableNumber(path=(2,)),
        "Hồ sơ chuyển tiền",
        headers=["STT", "Mục đích"],
    )

    assert out[0] == "2 Hồ sơ chuyển tiền  |  STT"
    assert not any("STT" in line for line in out[1:])


# ------------------------------------------------ the numbers Word draws itself
_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

_NUMBERING_XML = f"""<w:numbering {_W}>
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/></w:lvl>
    <w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="decimal"/></w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/></w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
    <w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>
  </w:num>
  <w:num w:numId="3"><w:abstractNumId w:val="1"/></w:num>
</w:numbering>"""


def _numbered_paragraph(num_id, ilvl=0):
    return parse_xml(
        f'<w:p {_W}><w:pPr><w:numPr>'
        f'<w:ilvl w:val="{ilvl}"/><w:numId w:val="{num_id}"/>'
        f"</w:numPr></w:pPr></w:p>"
    )


def _lists():
    return ListNumbering(parse_xml(_NUMBERING_XML))


def test_word_list_counters_are_replayed():
    numbering = _lists()
    seen = [numbering.advance(_numbered_paragraph(1)) for _ in range(3)]
    assert seen == [(1,), (2,), (3,)]


def test_a_deeper_level_restarts_under_each_parent():
    numbering = _lists()
    numbering.advance(_numbered_paragraph(1))
    assert numbering.advance(_numbered_paragraph(1, ilvl=1)) == (1, 1)
    assert numbering.advance(_numbered_paragraph(1, ilvl=1)) == (1, 2)
    assert numbering.advance(_numbered_paragraph(1)) == (2,)
    assert numbering.advance(_numbered_paragraph(1, ilvl=1)) == (2, 1)


def test_a_bulleted_list_names_no_section():
    assert _lists().advance(_numbered_paragraph(3)) is None


def test_a_restarted_list_starts_from_one_again():
    """An appendix that numbers itself 1., 2., 3. again is a new `w:num`."""
    numbering = _lists()
    for _ in range(4):
        numbering.advance(_numbered_paragraph(1))
    assert numbering.advance(_numbered_paragraph(2)) == (1,)


def test_an_unnumbered_paragraph_is_not_counted():
    assert _lists().advance(parse_xml(f"<w:p {_W}/>")) is None


# -------------------------------------------------------------- the Word path
def _docx_with(tmp_path, build, numbering=True):
    doc = Document()
    build(doc)
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc.save(src)
    DocxTableFlattener(verify_output=False, numbering=numbering).process(
        str(src), str(out)
    )
    return [p.text for p in Document(str(out)).paragraphs if p.text.strip()]


def _grid_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    return table


def test_word_table_is_numbered_under_the_numbered_paragraph_above_it(tmp_path):
    def build(doc):
        doc.add_paragraph("1. Thuật ngữ")
        _grid_table(doc, [["Tên", "Tuổi"], ["Nam", "25"], ["Lan", "30"]])

    texts = _docx_with(tmp_path, build)

    # Section 1 holds nothing but this table, so the table is section 1 -- the
    # heading already standing above it is its name, and no line repeats it.
    assert texts == [
        "1. Thuật ngữ",
        "1.1 Tên: Nam  |  Tuổi: 25",
        "1.2 Tên: Lan  |  Tuổi: 30",
    ]


def test_the_number_word_draws_decides_where_the_table_sits(tmp_path):
    """The section number is nowhere in the text -- Word paints it from a list.

    This is the common shape of an official document: the author ticked
    "numbered list", so the paragraph that reads "2. Chính sách ưu đãi" on
    screen reaches us as "Chính sách ưu đãi".
    """
    def build(doc):
        doc.add_paragraph("Đối tượng áp dụng", style="List Number")
        doc.add_paragraph("Chính sách ưu đãi", style="List Number")
        _grid_table(doc, [["Hạng", "Ưu đãi"], ["Gold", "0,10"]])

    texts = _docx_with(tmp_path, build)

    assert texts[-1] == "2.1 Hạng: Gold  |  Ưu đãi: 0,10"


def test_a_numbered_paragraph_inside_a_table_still_counts(tmp_path):
    """Word counts it, so the section after the table must not be one short."""
    def build(doc):
        doc.add_paragraph("Mục một", style="List Number")
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).paragraphs[0].text = "Mục hai nằm trong bảng"
        table.cell(0, 0).paragraphs[0].style = "List Number"
        doc.add_paragraph("Mục ba", style="List Number")
        _grid_table(doc, [["Hạng", "Ưu đãi"], ["Gold", "0,10"]])

    texts = _docx_with(tmp_path, build)

    assert any(t.startswith("3.1 Hạng: Gold") for t in texts), texts


def test_word_heading_styles_number_themselves(tmp_path):
    def build(doc):
        doc.add_heading("Quy định chung", level=1)
        doc.add_heading("Biểu phí", level=2)
        _grid_table(doc, [["Tuổi", "Phí"], ["30", "1.000"]])

    texts = _docx_with(tmp_path, build)

    # "Biểu phí" is section 1.1 and holds only this table, so the rows are 1.1.1
    # onwards -- three components, still inside the numbering limit.
    assert texts == ["Quy định chung", "Biểu phí", "1.1.1 Tuổi: 30  |  Phí: 1.000"]


def test_a_second_table_in_the_same_section_gets_the_next_number(tmp_path):
    def build(doc):
        doc.add_paragraph("2. Quyền lợi")
        _grid_table(doc, [["A", "B"], ["1", "2"]])
        doc.add_paragraph("Đoạn văn giữa hai bảng.")
        _grid_table(doc, [["C", "D"], ["3", "4"]])

    texts = _docx_with(tmp_path, build)

    assert "2.1.1 A: 1  |  B: 2" in texts
    assert "2.2.1 C: 3  |  D: 4" in texts


def test_a_table_the_paragraph_above_does_not_name_gets_no_caption(tmp_path):
    """Prose in front of a table says nothing about it, so nothing is invented."""
    def build(doc):
        doc.add_paragraph("2. Quyền lợi")
        doc.add_paragraph("Đoạn văn trước bảng, không phải tên bảng.")
        _grid_table(doc, [["Hạng", "Ưu đãi"], ["Gold", "0,10"]])

    texts = _docx_with(tmp_path, build)

    # Nothing was inserted between the prose and the rows.
    assert texts == [
        "2. Quyền lợi",
        "Đoạn văn trước bảng, không phải tên bảng.",
        "2.1 Hạng: Gold  |  Ưu đãi: 0,10",
    ]


def test_a_lead_in_sentence_names_the_table_it_hands_over_to(tmp_path):
    """Two tables in one section, so each still opens a level and is captioned."""
    def build(doc):
        doc.add_paragraph("2. Quyền lợi")
        doc.add_paragraph("Chức danh lãnh đạo và phân hạng tương ứng như bảng sau:")
        _grid_table(doc, [["Hạng", "Ưu đãi"], ["Gold", "0,10"]])
        doc.add_paragraph("Hạn mức tương ứng như bảng sau:")
        _grid_table(doc, [["Hạng", "Hạn mức"], ["Gold", "500"]])

    texts = _docx_with(tmp_path, build)

    assert "2.1 Chức danh lãnh đạo và phân hạng tương ứng như bảng sau" in texts
    assert "2.2 Hạn mức tương ứng như bảng sau" in texts


def test_a_section_holding_one_table_does_not_caption_it_a_second_time(tmp_path):
    """The heading above already names the table, under the very same number."""
    def build(doc):
        doc.add_paragraph("3.5. Quy định sản phẩm tiết kiệm thường")
        _grid_table(doc, [["Kỳ hạn", "Lãi suất"], ["1 tháng", "3,5%"]])

    texts = _docx_with(tmp_path, build)

    assert texts == [
        "3.5. Quy định sản phẩm tiết kiệm thường",
        "3.5.1 Kỳ hạn: 1 tháng  |  Lãi suất: 3,5%",
    ]


def test_numbering_loses_nothing_of_a_word_document(tmp_path):
    """Numbering only ever puts a number in front of a line."""
    doc = Document()
    doc.add_paragraph("1. Thuật ngữ")
    _grid_table(doc, [["Khoản", "Quy định"], ["3.8", "Giải ngân"]])
    doc.add_paragraph("Đoạn văn giữa hai bảng.")
    _grid_table(doc, [["Mã", "Tên"], ["A1", "Bút bi"]])
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc.save(src)

    summary = DocxTableFlattener().process(str(src), str(out))

    assert summary["total_tables_flattened"] == 2
    assert summary["verification"].passed, summary["verification"].describe()


def test_turning_numbering_off_restores_the_plain_bullets(tmp_path):
    def build(doc):
        doc.add_paragraph("1. Thuật ngữ")
        _grid_table(doc, [["Tên", "Tuổi"], ["Nam", "25"]])

    assert _docx_with(tmp_path, build, numbering=False) == [
        "1. Thuật ngữ",
        "- Tên: Nam  |  Tuổi: 25",
    ]


def test_a_nested_table_stays_part_of_its_parent_row(tmp_path):
    """A sub-table is content of a cell, not a section of the document."""
    def build(doc):
        doc.add_paragraph("1. Hạn mức")
        table = _grid_table(doc, [["Điều kiện", "Nội dung"], ["Hạn mức", ""]])
        nested = table.cell(1, 1).add_table(rows=2, cols=2)
        nested.style = "Table Grid"
        for r, row in enumerate([["Chức danh", "HMTC"], ["Giám đốc", "500"]]):
            for c, text in enumerate(row):
                nested.cell(r, c).text = text

    texts = _docx_with(tmp_path, build)

    assert texts[1].startswith("1.1 ")
    assert any("Chức danh: Giám đốc  |  HMTC: 500" in t for t in texts)
    # The sub-table did not take a number of its own from the outline.
    assert not any(t.startswith("1.2 ") for t in texts)


# ------------------------------------------------------------- the Excel path
def _xlsx_paragraphs(tmp_path, build, numbering=True):
    workbook = Workbook()
    build(workbook)
    src = tmp_path / "book.xlsx"
    out = tmp_path / "book.docx"
    workbook.save(src)
    XlsxTableFlattener(verify_output=False, numbering=numbering).process(
        str(src), str(out)
    )
    return [p.text for p in Document(str(out)).paragraphs if p.text.strip()]


def test_each_sheet_of_a_workbook_becomes_a_numbered_section(tmp_path):
    def build(workbook):
        first = workbook.active
        first.title = "Một"
        first.append(["Tên", "Tuổi"])
        first.append(["Nam", 25])
        second = workbook.create_sheet("Hai")
        second.append(["Mã", "Giá"])
        second.append(["X1", 10])

    # A sheet holding a single table is that table: its rows carry on from the
    # sheet's own number rather than opening a level under it.
    assert _xlsx_paragraphs(tmp_path, build) == [
        "1. Một",
        "1.1 Tên: Nam  |  Tuổi: 25",
        "2. Hai",
        "2.1 Mã: X1  |  Giá: 10",
    ]


def test_a_line_above_a_table_in_a_sheet_names_it(tmp_path):
    """A sheet has no headings: a line of its own is all a table can be named by."""
    def build(workbook):
        sheet = workbook.active
        sheet["A1"] = "Biểu phí bảo hiểm 2026"
        sheet.append([])
        sheet.append(["Tuổi", "Phí"])
        sheet.append([30, 1000])

    assert _xlsx_paragraphs(tmp_path, build) == [
        "- Biểu phí bảo hiểm 2026",
        "1 Biểu phí bảo hiểm 2026",
        "1.1 Tuổi: 30  |  Phí: 1000",
    ]


def test_a_single_sheet_numbers_its_tables_from_the_top(tmp_path):
    def build(workbook):
        sheet = workbook.active
        sheet.append(["Tên", "Tuổi"])
        sheet.append(["Nam", 25])
        sheet.append([])
        sheet.append(["Mã", "Giá"])
        sheet.append(["X1", 10])

    assert _xlsx_paragraphs(tmp_path, build) == [
        "1.1 Tên: Nam  |  Tuổi: 25",
        "2.1 Mã: X1  |  Giá: 10",
    ]


# --------------------------------------------------------------- the PDF path
def _write(page, items, size=10):
    font = fitz.Font(fontfile=settings.get_font_path())
    writer = fitz.TextWriter(page.rect)
    for x, y, text in items:
        writer.append(fitz.Point(x, y), text, font=font, fontsize=size)
    writer.write_text(page)


def _grid(page, x0, y0, col_w, row_h, ncols, nrows):
    for r in range(nrows + 1):
        page.draw_line(
            fitz.Point(x0, y0 + r * row_h),
            fitz.Point(x0 + ncols * col_w, y0 + r * row_h),
            width=0.6,
        )
    for c in range(ncols + 1):
        page.draw_line(
            fitz.Point(x0 + c * col_w, y0),
            fitz.Point(x0 + c * col_w, y0 + nrows * row_h),
            width=0.6,
        )


def test_a_pdf_table_is_numbered_under_the_heading_printed_above_it(tmp_path):
    src = tmp_path / "src.pdf"
    out = tmp_path / "out.pdf"
    doc = fitz.open()
    page = doc.new_page()
    _write(page, [(60, 70, "2. Quy dinh chung")])
    _write(page, [(60, 92, "2.1. Bieu phi bao hiem")])
    rows = [["Tuoi", "Phi"], ["30", "1.000"], ["40", "2.000"]]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            _write(page, [(70 + c * 150, 130 + r * 25, text)])
    _grid(page, 60, 110, 150, 25, 2, 3)
    doc.save(src)
    doc.close()

    summary = PDFTableFlattenerPipeline().process(str(src), str(out))
    assert summary["verification"].passed, summary["verification"].describe()

    with fitz.open(out) as flat:
        text = flat[0].get_text()

    # Section 2.1 holds nothing but this table, so the table is 2.1 and its rows
    # carry on from the heading printed above them.
    assert "2.1. Bieu phi bao hiem" in text
    assert "2.1.1 Tuoi: 30  |  Phi: 1.000" in text
    assert "2.1.2 Tuoi: 40  |  Phi: 2.000" in text
    # No caption repeating that heading a line below it, one level deeper.
    assert "2.1.1 Bieu phi bao hiem" not in text
    # The document's own headings are passthrough content and stay as they were.
    assert "2. Quy dinh chung" in text
    assert "2.1. Bieu phi bao hiem" in text
