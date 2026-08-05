"""The Word path: same bullets as the PDF path, written back into the .docx."""

import pytest

from docx import Document
from docx.shared import Pt

from pdf_table_tool.docx_flattener import DocxTableFlattener
from pdf_table_tool.verifier import _docx_table_count, _docx_tokens, verify_docx


def _grid_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    return table


def _run(tmp_path, doc, name="doc", numbering=False):
    src = tmp_path / f"{name}.docx"
    out = tmp_path / f"{name}_out.docx"
    doc.save(src)
    # These tests are about which cell is paired with which label, so they run
    # without the outline numbering that :mod:`test_numbering` covers -- a "1.1"
    # in front of every expected line would only make them harder to read.
    summary = DocxTableFlattener(numbering=numbering).process(str(src), str(out))
    return out, summary


def _paragraphs(path):
    return [p.text for p in Document(str(path)).paragraphs if p.text.strip()]


# ------------------------------------------------------------------- basics
def test_table_becomes_bullets_and_disappears(tmp_path):
    doc = Document()
    _grid_table(
        doc,
        [
            ["Tên", "Chức vụ"],
            ["Nguyễn Văn A", "Chuyên viên"],
            ["Trần Thị B", "Trưởng phòng"],
        ],
    )
    out, summary = _run(tmp_path, doc)

    assert summary["total_tables_flattened"] == 1
    assert _docx_table_count(str(out)) == 0
    assert _paragraphs(out) == [
        "- Tên: Nguyễn Văn A  |  Chức vụ: Chuyên viên",
        "- Tên: Trần Thị B  |  Chức vụ: Trưởng phòng",
    ]


def test_output_is_a_docx_that_still_opens(tmp_path):
    doc = Document()
    _grid_table(doc, [["Khoản", "Quy định"], ["3.1", "Nội dung"]])
    out, _ = _run(tmp_path, doc)

    assert out.suffix == ".docx"
    assert Document(str(out)) is not None


def test_content_around_the_table_is_untouched(tmp_path):
    doc = Document()
    doc.add_heading("Tiêu đề tài liệu", level=1)
    doc.add_paragraph("Đoạn văn phía trên bảng.")
    _grid_table(doc, [["Khoản", "Quy định"], ["3.1", "Giải ngân"]])
    doc.add_paragraph("Đoạn văn phía dưới bảng.")

    out, _ = _run(tmp_path, doc)
    texts = _paragraphs(out)

    assert texts[:2] == ["Tiêu đề tài liệu", "Đoạn văn phía trên bảng."]
    assert texts[-1] == "Đoạn văn phía dưới bảng."
    assert Document(str(out)).paragraphs[0].style.name.startswith("Heading")


# ------------------------------------------------------------------- layout
def test_vertically_merged_label_applies_to_every_row(tmp_path):
    doc = Document()
    table = _grid_table(
        doc,
        [
            ["Nhóm", "Hạng", "Hạn mức"],
            ["", "I", "1.000"],
            ["", "II", "700"],
        ],
    )
    table.cell(1, 0).merge(table.cell(2, 0)).text = "Nhóm A"

    out, _ = _run(tmp_path, doc)
    texts = _paragraphs(out)

    assert texts == [
        "- Nhóm A  |  Hạng: I  |  Hạn mức: 1.000",
        "- Nhóm A  |  Hạng: II  |  Hạn mức: 700",
    ]


def test_merged_header_labels_every_column_it_covers(tmp_path):
    """Word says the cell was merged on purpose, so it names both columns."""
    doc = Document()
    table = _grid_table(
        doc,
        [
            ["Năm", "", ""],
            ["1", "50%", "1,5%"],
            ["2", "30%", "1,5%"],
        ],
    )
    table.cell(0, 1).merge(table.cell(0, 2)).text = "Tỷ lệ"

    out, _ = _run(tmp_path, doc)

    assert _paragraphs(out) == [
        "- Năm: 1  |  Tỷ lệ: 50%  |  Tỷ lệ: 1,5%",
        "- Năm: 2  |  Tỷ lệ: 30%  |  Tỷ lệ: 1,5%",
    ]


def test_two_row_header_band_joins_group_and_sub_label(tmp_path):
    """A merged group label over sub-labels is one header two rows tall."""
    doc = Document()
    table = _grid_table(
        doc,
        [
            ["Năm", "", ""],
            ["", "Cơ bản", "Đóng thêm"],
            ["1", "50%", "1,5%"],
        ],
    )
    table.cell(0, 1).merge(table.cell(0, 2)).text = "Tỷ lệ"
    table.cell(0, 0).merge(table.cell(1, 0)).text = "Năm"

    out, _ = _run(tmp_path, doc)

    assert _paragraphs(out) == [
        "- Năm: 1  |  Tỷ lệ Cơ bản: 50%  |  Tỷ lệ Đóng thêm: 1,5%"
    ]


def test_a_sentence_merged_across_the_top_is_not_a_header(tmp_path):
    """The Word pass must not turn a two-dimensional table upside down.

    Labels run down column 0 and the merged top cell is a value, not a column
    name -- so the table still pivots into one bullet per case.
    """
    doc = Document()
    table = _grid_table(
        doc,
        [
            ["Điều kiện", "", ""],
            ["Tình huống", "Thỏa điều kiện", "Không thỏa điều kiện"],
            ["Thứ tự phân bổ phí", "Đóng cho phí định kỳ", "Đóng cho phí cơ bản"],
        ],
    )
    table.cell(0, 1).merge(table.cell(0, 2)).text = (
        "Khoản phí đóng vào có thể đủ để duy trì hiệu lực của Hợp đồng bảo hiểm"
    )

    out, _ = _run(tmp_path, doc)
    texts = _paragraphs(out)

    assert len(texts) == 2
    assert texts[0].startswith("- Điều kiện: Khoản phí đóng vào")
    assert "Tình huống: Thỏa điều kiện" in texts[0]
    assert "Tình huống: Không thỏa điều kiện" in texts[1]


def test_repeat_as_header_row_is_believed(tmp_path):
    """`Repeat as header row` is Word stating outright which row is the header.

    Geometry alone rejects this one: the first row fills fewer cells than the
    rows below it, which on a PDF means a cell continued from the page before.
    """
    doc = Document()
    table = _grid_table(
        doc,
        [
            ["Khoản", "Quy định", ""],
            ["3.8", "Giải ngân", "Chuyển khoản"],
            ["3.9", "Trả nợ", "Hàng tháng"],
        ],
    )
    table.rows[0]._tr.get_or_add_trPr().append(
        table.rows[0]._tr.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader",
            {},
        )
    )

    out, _ = _run(tmp_path, doc)

    assert _paragraphs(out) == [
        "- Khoản: 3.8  |  Quy định: Giải ngân  |  Chuyển khoản",
        "- Khoản: 3.9  |  Quy định: Trả nợ  |  Hàng tháng",
    ]


def test_multi_paragraph_cell_becomes_indented_sub_bullets(tmp_path):
    doc = Document()
    table = _grid_table(doc, [["Điều kiện", "Nội dung"], ["Vay vốn", ""]])
    cell = table.cell(1, 1)
    cell.text = "Khách hàng phải đáp ứng:"
    cell.add_paragraph("Có tài sản bảo đảm", style="List Bullet")
    cell.add_paragraph("Không có nợ xấu", style="List Bullet")

    out, _ = _run(tmp_path, doc)
    document = Document(str(out))
    lines = [p for p in document.paragraphs if p.text.strip()]

    assert lines[0].text == "- Điều kiện: Vay vốn  |  Nội dung:"
    assert [p.text for p in lines[1:]] == [
        "+ Khách hàng phải đáp ứng:",
        "- Có tài sản bảo đảm",
        "- Không có nợ xấu",
    ]
    # The sub-bullets are really indented, not merely prefixed.
    assert lines[0].paragraph_format.left_indent in (None, 0)
    assert all(p.paragraph_format.left_indent > 0 for p in lines[1:])


def test_nested_table_keeps_its_own_column_pairing(tmp_path):
    doc = Document()
    table = _grid_table(doc, [["Điều kiện", "Nội dung"], ["Hạn mức", ""]])
    nested = table.cell(1, 1).add_table(rows=2, cols=2)
    nested.style = "Table Grid"
    for r, row in enumerate([["Nhóm chức danh", "HMTC"], ["Giám đốc", "500 triệu"]]):
        for c, text in enumerate(row):
            nested.cell(r, c).text = text

    out, _ = _run(tmp_path, doc)
    texts = _paragraphs(out)

    assert _docx_table_count(str(out)) == 0
    # The sub-table's own pairing survives: its two columns stay joined to each
    # other instead of dissolving into loose words of the parent cell.
    assert any(
        "Nhóm chức danh: Giám đốc  |  HMTC: 500 triệu" in t for t in texts
    ), texts


def test_manual_line_break_starts_a_new_bullet(tmp_path):
    doc = Document()
    table = _grid_table(doc, [["Khoản", "Quy định"], ["3.1", ""]])
    paragraph = table.cell(1, 1).paragraphs[0]
    run = paragraph.add_run("Dòng một")
    run.add_break()
    paragraph.add_run("Dòng hai")

    out, _ = _run(tmp_path, doc)
    texts = _paragraphs(out)

    assert any("Dòng một" in t for t in texts)
    assert any("Dòng hai" in t for t in texts)
    assert not any("Dòng mộtDòng hai" in t for t in texts)


# -------------------------------------------------------------- other stories
def test_table_in_a_page_header_is_flattened_too(tmp_path):
    doc = Document()
    doc.add_paragraph("Thân tài liệu.")
    header = doc.sections[0].header
    table = header.add_table(rows=1, cols=2, width=Pt(400))
    table.cell(0, 0).text = "Mã biểu"
    table.cell(0, 1).text = "BM-01"

    out, summary = _run(tmp_path, doc)

    assert summary["total_tables_flattened"] == 1
    assert _docx_table_count(str(out)) == 0
    assert "- Mã biểu  |  BM-01" in [
        p.text for p in Document(str(out)).sections[0].header.paragraphs
    ]


# ------------------------------------------------------------- verification
def test_no_token_of_the_source_is_lost(tmp_path):
    doc = Document()
    doc.add_paragraph("Đoạn văn ngoài bảng.")
    _grid_table(
        doc,
        [
            ["Khoản", "Quy định"],
            ["3.8", "Giải ngân bằng chuyển khoản"],
            ["3.9", "Trả nợ hàng tháng"],
        ],
    )
    out, summary = _run(tmp_path, doc)

    report = summary["verification"]
    assert report.passed, report.describe()
    assert not (_docx_tokens(str(tmp_path / "doc.docx")) - _docx_tokens(str(out)))


def test_verification_reports_a_table_left_behind(tmp_path):
    """The criterion has teeth: an untouched table must fail the check."""
    doc = Document()
    _grid_table(doc, [["Khoản", "Quy định"], ["3.1", "Nội dung"]])
    src = tmp_path / "same.docx"
    doc.save(src)

    report = verify_docx(str(src), str(src), [])
    assert not report.criterion_2_ok
    assert report.criterion_1_and_2_ok  # nothing lost -- only the table remains


def test_empty_table_leaves_no_stray_paragraph_text(tmp_path):
    doc = Document()
    doc.add_paragraph("Trước.")
    _grid_table(doc, [["", ""], ["", ""]])
    doc.add_paragraph("Sau.")

    out, summary = _run(tmp_path, doc)

    assert summary["total_tables_flattened"] == 0
    assert _docx_table_count(str(out)) == 0
    assert _paragraphs(out) == ["Trước.", "Sau."]


@pytest.mark.parametrize("style", ["List Bullet", "List Number"])
def test_list_paragraphs_keep_their_glyph_out_of_the_text(tmp_path, style):
    doc = Document()
    table = _grid_table(doc, [["Điều kiện", "Nội dung"], ["Vay vốn", ""]])
    cell = table.cell(1, 1)
    cell.text = "Yêu cầu:"
    cell.add_paragraph("Có tài sản bảo đảm", style=style)

    out, _ = _run(tmp_path, doc, name=style.replace(" ", "_"))
    texts = _paragraphs(out)

    assert any(t.endswith("Có tài sản bảo đảm") for t in texts)
