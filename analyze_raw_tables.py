import os
import sys
import docx
from doc_table_converter import convert_doc_or_pdf_to_docx, clean_table_grid, is_header_footer_table

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

pdf_files = [
    "Bảo hiểm nhân thọ - Lộc Phát Hưng Thịnh - Quy định sản phẩm.pdf",
    "Bảo hiểm nhân thọ - Lộc Phát Tràng An - Quy định sản phẩm.pdf",
    "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf",
    "Sản phẩm cho vay kinh doanh xi măng Xuân Thành - Kênh quầy.docx.pdf",
    "Sản phẩm cho vay tái tài trợ - kênh quầy.docx.pdf"
]

for idx, pdf_name in enumerate(pdf_files, 1):
    print(f"\n==================================================")
    print(f"ANALYZING ORIGINAL TABLES IN FILE {idx}: {pdf_name}")
    print(f"==================================================")
    docx_path, is_tmp = convert_doc_or_pdf_to_docx(pdf_name)
    doc = docx.Document(docx_path)
    
    print(f"Total tables: {len(doc.tables)}")
    
    for t_i, table in enumerate(doc.tables, 1):
        grid = clean_table_grid(table)
        if is_header_footer_table(grid):
            continue
        print(f"\n--- Table {t_i} (Rows: {len(grid)}, Cols: {max(len(r) for r in grid) if grid else 0}) ---")
        for r_i, row in enumerate(grid):
            print(f"  Row {r_i}: {row}")
            
    if is_tmp and os.path.exists(docx_path):
        os.remove(docx_path)
