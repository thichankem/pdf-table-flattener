import os
import sys
import docx
from pdf2docx import Converter

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

pdf_name = "Bảo hiểm nhân thọ - Lộc Phát Tràng An - Quy định sản phẩm.pdf"
temp_docx = "temp_inspect_img.docx"

cv = Converter(pdf_name)
cv.convert(temp_docx)
cv.close()

doc = docx.Document(temp_docx)

print("=== INSPECTING ALL DRAWINGS IN FILE 2 ===")

# 1. Inspect drawings in main body paragraphs vs tables
drawings_in_body_paras = 0
drawings_in_tables = 0

for p_i, p in enumerate(doc.paragraphs):
    if 'w:drawing' in p._element.xml or 'w:pict' in p._element.xml:
        drawings_in_body_paras += 1
        print(f"Paragraph {p_i} has drawing! Text: '{p.text[:40]}'")

for t_i, table in enumerate(doc.tables):
    if 'w:drawing' in table._element.xml or 'w:pict' in table._element.xml:
        drawings_in_tables += 1
        print(f"Table {t_i} has drawing! Rows: {len(table.rows)}, Cols: {len(table.columns)}")
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                if 'w:drawing' in cell._element.xml or 'w:pict' in cell._element.xml:
                    print(f"   Cell ({r_i},{c_i}) has drawing! Text: '{cell.text.strip()[:40]}'")

print(f"\nTotal drawings in body paragraphs: {drawings_in_body_paras}")
print(f"Total drawings in tables: {drawings_in_tables}")

if os.path.exists(temp_docx):
    os.remove(temp_docx)
