import os
import sys
import re
import tempfile
from pathlib import Path
import docx
from docx.document import Document as DocumentClass
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt

def convert_doc_or_pdf_to_docx(file_path: str) -> tuple[str, bool]:
    """
    Converts .doc (via MS Word COM) or .pdf (via pdf2docx) to .docx.
    Returns (docx_path, is_temporary).
    """
    abs_path = os.path.abspath(file_path)
    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

    ext = os.path.splitext(abs_path)[1].lower()

    if ext == ".docx":
        return abs_path, False

    temp_dir = tempfile.gettempdir()
    temp_docx = os.path.join(temp_dir, f"conv_{os.path.basename(abs_path)}.docx")

    if ext == ".doc":
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(abs_path)
                doc.SaveAs2(temp_docx, FileFormat=16)
                doc.Close()
                return temp_docx, True
            finally:
                word.Quit()
        except Exception as e:
            raise RuntimeError(f"Lỗi khi chuyển đổi file .doc bằng Word: {e}")

    elif ext == ".pdf":
        try:
            from pdf2docx import Converter
            cv = Converter(abs_path)
            cv.convert(temp_docx)
            cv.close()
            return temp_docx, True
        except Exception as e:
            raise RuntimeError(f"Lỗi khi chuyển đổi PDF sang DOCX: {e}")

    return abs_path, False

def get_table_text_grid(table: Table) -> list[list[str]]:
    """
    Extracts 2D array of text strings from docx Table, cleaning up newlines and merged cells.
    """
    grid = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            raw_text = cell.text.strip()
            clean_text = " ".join(raw_text.splitlines())
            row_cells.append(clean_text)
        grid.append(row_cells)
    return grid

def clean_text_string(text: str) -> str:
    """
    Cleans non-breaking space, zero-width space, BOM, null bytes, form feeds, and excessive spaces.
    """
    if not text:
        return ""
    t = text.replace('\xa0', ' ').replace('\u200b', '').replace('\ufeff', '').replace('\x00', '').replace('\x0c', '')
    lines = [re.sub(r'[ \t]+', ' ', l).strip() for l in t.splitlines() if l.strip()]
    return " ".join(lines).strip()

def get_table_text_grid(table: Table) -> list[list[str]]:
    """
    Extracts 2D array of text strings from docx Table, cleaning up newlines and merged cells.
    """
    grid = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            raw_text = cell.text
            clean_text = clean_text_string(raw_text)
            row_cells.append(clean_text)
        grid.append(row_cells)
    return grid

def is_bullet_marker(text: str) -> bool:
    """
    Checks if a text string is a bullet/numbered list marker or dash/punct rather than a table value/header.
    Plain numbers like '0', '1', '2' inside data cells are NOT bullet markers.
    Only numbers with dots/parens like '1.', '(1)' are list markers.
    """
    s = text.strip().lower()
    if not s or s in ['•', '▪', '*', '-', '+', '–', '—', '(v)', '(i)', '(ii)', '(iii)', '(iv)', 'a.', 'b.', 'c.', 'd.']:
        return True
    if re.match(r'^[-—–\.\*\•\+\s]+$', s):
        return True
    if re.match(r'^(\([a-z0-9ivxlcdm]+\)|[a-z0-9]+\.|\([0-9]+\))$', s, re.IGNORECASE):
        return True
    return False

def clean_table_grid(table: Table) -> list[list[str]]:
    """
    Extracts 2D grid from Table, cleaning newlines and filtering out page header/footer rows.
    """
    grid = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            clean_text = clean_text_string(cell.text)
            row_cells.append(clean_text)
        grid.append(row_cells)

    # Filter out page header/footer rows (containing page numbers like X/34 or running title)
    cleaned_grid = []
    for row in grid:
        row_str = " ".join(row).strip()
        if re.search(r'\b\d+/\d+\b', row_str) and ('Sản phẩm bảo hiểm' in row_str or 'Dai-ichi' in row_str or len(row_str) < 150):
            continue
        if not any(c.strip() for c in row):
            continue
        cleaned_grid.append(row)

    return cleaned_grid

def is_header_footer_table(grid: list[list[str]]) -> bool:
    """
    Checks if table is just a PDF page header/footer box rather than a real data table.
    """
    if not grid or len(grid) == 0:
        return True
    if len(grid) == 1:
        row_str = " ".join(grid[0]).strip()
        if re.search(r'\b\d+/\d+\b', row_str) or ('Sản phẩm bảo hiểm' in row_str and len(row_str) < 150):
            return True
    return False

def is_likely_header_row(row: list[str]) -> bool:
    """
    Determines if row 0 of a grid is a genuine column header row vs data row.
    """
    if not row or not any(row):
        return False
    # Allow descriptive Vietnamese column header cells up to 250 chars
    if any(len(c) > 250 for c in row if c):
        return False
    first = row[0].strip()
    # If first cell is a section code like 3.1, 3.2, 3.16
    if re.match(r'^\d+(\.\d+)+\.?$', first):
        return False
    # If first cell is a bullet marker
    if is_bullet_marker(first):
        return False
    return True

def format_table_to_dash_text(grid: list[list[str]], separator: str = " | ", use_header: bool = True, show_row_indices: bool = False, bullet_prefix: bool = True) -> str:
    """
    Converts 2D data table grid to formatted key-value bullet text lines (- Tên: Nam | Tuổi: 25 | Chức vụ: Dev).
    Intelligently transposes 2-row horizontal matrix tables and handles Vertical Scenario Matrix tables
    (e.g. Điều kiện / Tình huống / Thứ tự phân bổ phí) semantically without losing column associations.
    Never injects fake 'Cột 1', 'Cột 2' labels or empty cell key names.
    """
    if not grid or not any(grid):
        return ""

    num_rows = len(grid)
    num_cols = max(len(r) for r in grid)
    lines = []
    prefix_str = "- " if bullet_prefix else ""

    # 1. Special handling for 2-row horizontal matrix tables (Row 0 metric vs Row 1 metric across columns 1..N)
    if num_rows == 2 and num_cols > 1:
        metric0 = clean_text_string(grid[0][0]) if len(grid[0]) > 0 else ""
        metric1 = clean_text_string(grid[1][0]) if len(grid[1]) > 0 else ""
        
        if metric0 and metric1 and not is_bullet_marker(metric0):
            for c in range(1, num_cols):
                val0 = clean_text_string(grid[0][c]) if c < len(grid[0]) else ""
                val1 = clean_text_string(grid[1][c]) if c < len(grid[1]) else ""
                if not val0 and not val1:
                    continue
                    
                h0 = f"{metric0}: {val0}" if val0 else ""
                h1 = f"{metric1}: {val1}" if val1 else ""
                
                parts = [p for p in [h0, h1] if p]
                if parts:
                    line = separator.join(parts)
                    lines.append(prefix_str + line)
            if lines:
                return "\n".join(lines).strip()

    # 2. Special handling for Vertical Scenario Matrix Tables (Col 0 contains row attribute keys across N scenario cols)
    # Key insight: If ANY row has merged cells (identical values across cols 1..N), this is a vertical matrix
    # even if is_likely_header_row returns True. Merged cells are the definitive signal for scenario matrices.
    # Standard header tables (e.g. Fund table) have DIFFERENT values in each header column.
    has_any_merged_row = False
    if num_cols > 2 and num_rows > 1:
        for r in grid:
            if len(r) > 2:
                val1 = clean_text_string(r[1]) if len(r) > 1 else ""
                if val1:  # Only check non-empty rows
                    all_same = True
                    for c_i in range(2, min(len(r), num_cols)):
                        val_c = clean_text_string(r[c_i])
                        if val_c != val1:
                            all_same = False
                            break
                    if all_same:
                        has_any_merged_row = True
                        break

    is_header_r0 = is_likely_header_row(grid[0])
    col0_is_labels = (num_cols > 1 and num_rows > 1 and (not is_header_r0 or has_any_merged_row))
    if col0_is_labels:
        for r in grid:
            val0 = clean_text_string(r[0]) if len(r) > 0 else ""
            if not val0 or is_bullet_marker(val0) or len(val0) > 80:
                col0_is_labels = False
                break

    if col0_is_labels:
        common_rows = []
        varying_rows = []
        
        for r_i, r in enumerate(grid):
            val1 = clean_text_string(r[1]) if len(r) > 1 else ""
            all_same = True
            for c_i in range(2, num_cols):
                val_c = clean_text_string(r[c_i]) if c_i < len(r) else ""
                if val_c != val1:
                    all_same = False
                    break
            if all_same:
                common_rows.append(r_i)
            else:
                varying_rows.append(r_i)

        if varying_rows:
            # Output common merged rows first
            for r_i in common_rows:
                k = clean_text_string(grid[r_i][0])
                v = clean_text_string(grid[r_i][1]) if len(grid[r_i]) > 1 else ""
                if k and v:
                    lines.append(f"{prefix_str}{k}: {v}")
                elif v:
                    lines.append(f"{prefix_str}{v}")
                    
            # Output scenario columns 1..N
            for c_i in range(1, num_cols):
                col_parts = []
                seen_col_parts = set()
                for r_i in varying_rows:
                    k = clean_text_string(grid[r_i][0])
                    v = clean_text_string(grid[r_i][c_i]) if c_i < len(grid[r_i]) else ""
                    if k and v:
                        part = f"{k}: {v}"
                    elif v:
                        part = v
                    else:
                        continue
                    if part not in seen_col_parts:
                        seen_col_parts.add(part)
                        col_parts.append(part)
                if col_parts:
                    lines.append(f"{prefix_str}{separator.join(col_parts)}")
            if lines:
                return "\n".join(lines).strip()

    # 3. Standard Table with Column Headers or Data Rows
    has_header = False
    headers = []
    data_rows = []

    if use_header and num_rows > 1 and is_likely_header_row(grid[0]):
        has_header = True
        raw_headers = grid[0]
        data_rows = grid[1:]
        for col_idx in range(num_cols):
            h_text = clean_text_string(raw_headers[col_idx]) if col_idx < len(raw_headers) else ""
            if is_bullet_marker(h_text) or h_text in ['-', '--', '---']:
                h_text = ""
            headers.append(h_text)
    else:
        data_rows = grid

    for idx, row in enumerate(data_rows, 1):
        row_parts = []
        seen_parts = set()
        for col_idx in range(num_cols):
            c_val = clean_text_string(row[col_idx]) if col_idx < len(row) else ""
            if not c_val:
                continue

            h_name = headers[col_idx] if (has_header and col_idx < len(headers)) else ""

            # If no header exists and c_val is just a bullet/dash marker like '-' or '•', skip it
            if not h_name and is_bullet_marker(c_val):
                continue

            if h_name:
                part = f"{h_name}: {c_val}"
            else:
                part = c_val

            if part in seen_parts:
                continue
            seen_parts.add(part)
            row_parts.append(part)

        if row_parts:
            line = separator.join(row_parts)
            if show_row_indices:
                lines.append(f"--- Dòng {idx} ---")
            lines.append(prefix_str + line)

    return "\n".join(lines).strip()

def fix_toc_paragraphs(doc: docx.Document):
    """
    Splits concatenated Table of Contents (Mục lục) paragraphs into individual clean lines per entry.
    """
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if ('MỤC LỤC' in t.upper() or re.search(r'\.{3,}\s*\d+', t)) and ('ĐIỀU 1:' in t or 'PHẦN 1:' in t):
            pattern = r'(\.{2,}\s*\d+|\b\d{1,3}\b)(?=\s*(?:ĐIỀU\s+\d+|PHẦN\s+\d+|\d+\.\d+\.|\d+\.\d+))'
            cleaned_text = re.sub(pattern, r'\1\n', t)
            lines = [l.strip() for l in cleaned_text.splitlines() if l.strip()]
            if len(lines) > 1:
                p.text = lines[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                parent_elm = p._element.getparent()
                p_idx = parent_elm.index(p._element)
                for l_offset, l_str in enumerate(lines[1:], 1):
                    np = doc.add_paragraph()
                    np.text = l_str
                    np.paragraph_format.space_before = Pt(0)
                    np.paragraph_format.space_after = Pt(2)
                    parent_elm.insert(p_idx + l_offset, np._element)

def unpack_text_table_inplace(doc: docx.Document, table: Table, grid: list[list[str]]):
    """
    Unpacks single-column callouts or 1-row list items back into normal paragraphs
    WITHOUT adding fake column headers or bullets.
    """
    tbl_elm = table._element
    parent_elm = tbl_elm.getparent()
    tbl_index = parent_elm.index(tbl_elm)

    lines = []
    for row in grid:
        for val in row:
            if val.strip():
                lines.append(val.strip())

    for offset, line_str in enumerate(lines):
        p = doc.add_paragraph()
        p.text = line_str
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        parent_elm.insert(tbl_index + offset, p._element)

    parent_elm.remove(tbl_elm)

def replace_table_inplace(doc: docx.Document, table: Table, text_content: str):
    """
    Replaces a Table element IN-PLACE in the original document XML tree
    with compact new Paragraph elements for each text line, deleting the table.
    """
    tbl_elm = table._element
    parent_elm = tbl_elm.getparent()
    tbl_index = parent_elm.index(tbl_elm)

    lines = text_content.splitlines()
    for offset, line_str in enumerate(lines):
        p = doc.add_paragraph()
        p.text = line_str
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

        p_elm = p._element
        parent_elm.insert(tbl_index + offset, p_elm)

    parent_elm.remove(tbl_elm)

def has_drawing(p: Paragraph) -> bool:
    """
    Checks if a paragraph contains image/drawing/shape XML elements.
    """
    xml = p._element.xml
    return ('w:drawing' in xml) or ('w:pict' in xml)

def cleanup_document_blank_spaces(doc: docx.Document):
    """
    Scans document to eliminate consecutive empty paragraphs, clean strange/invisible characters,
    and normalize excessive paragraph spacing while PRESERVING 100% of images and drawings.
    """
    paragraphs = list(doc.paragraphs)
    consecutive_empty = 0

    for p in paragraphs:
        if has_drawing(p):
            consecutive_empty = 0
            continue

        raw = p.text
        cleaned = clean_text_string(raw)
        if cleaned != raw:
            p.text = cleaned

        if not cleaned:
            consecutive_empty += 1
            if consecutive_empty > 1:
                p_elm = p._element
                if p_elm.getparent() is not None:
                    p_elm.getparent().remove(p_elm)
        else:
            consecutive_empty = 0
            if p.paragraph_format.space_after and p.paragraph_format.space_after.pt > 6:
                p.paragraph_format.space_after = Pt(4)

def is_footer_or_header_page(text: str) -> bool:
    """
    Checks if line is header/footer page marker like 'Page 1' or 'Document Title \t 1/4'.
    """
    s = text.strip()
    if re.search(r'\t\d+/\d+$', s) or re.search(r'^\s*page\s+\d+', s, re.I):
        return True
    return False

def process_pseudo_tables(doc: docx.Document, separator: str = " | ", use_header: bool = True, show_row_indices: bool = False, bullet_prefix: bool = True) -> list[str]:
    """
    Identifies tab-delimited multi-column pseudo tables in document paragraphs,
    formats them into key-value text lines, replaces paragraphs in-place, and returns extracted text blocks.
    """
    paragraphs = list(doc.paragraphs)
    extracted_texts = []
    in_toc_section = False
    i = 0

    while i < len(paragraphs):
        p = paragraphs[i]
        p_text = p.text.strip()

        if re.search(r'^\s*mục\s+lục\s*$', p_text, re.IGNORECASE):
            in_toc_section = True
            i += 1
            continue
        elif in_toc_section and re.search(r'^\s*(phần|điều)\s+\d+', p_text, re.IGNORECASE):
            in_toc_section = False

        if in_toc_section or is_toc_paragraph(p) or is_footer_or_header_page(p.text) or not p_text:
            i += 1
            continue

        if '\t' in p.text:
            cols = [c.strip() for c in p.text.split('\t') if c.strip()]
            is_start = False
            if len(cols) >= 3:
                is_start = True
            elif len(cols) == 2 and not is_bullet_marker(cols[0]):
                is_start = True

            if is_start:
                grid = []
                consumed_paras = []

                grid.append(list(cols))
                consumed_paras.append(p)
                i += 1

                while i < len(paragraphs):
                    p_next = paragraphs[i]
                    t_next = p_next.text.strip()

                    if is_footer_or_header_page(p_next.text):
                        i += 1
                        continue

                    if not t_next:
                        i += 1
                        continue

                    if '\t' in p_next.text:
                        next_cols = [c.strip() for c in p_next.text.split('\t') if c.strip()]
                        is_next_row = False
                        if len(next_cols) >= 3:
                            is_next_row = True
                        elif len(next_cols) == 2 and not is_bullet_marker(next_cols[0]):
                            is_next_row = True

                        if is_next_row:
                            grid.append(list(next_cols))
                            consumed_paras.append(p_next)
                            i += 1
                            continue

                    if re.match(r'^\d+\.\s+[A-ZĐÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬEÈÉẺẼẸÊẾỀỂỄỆIÌÍỈĨỊOÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢUÙÚỦŨỤƯỨỪỬỮỰYỲÝỶỸỴ]', t_next):
                        break

                    if grid and len(grid[-1]) > 0:
                        grid[-1][-1] += " " + t_next
                        consumed_paras.append(p_next)
                    i += 1

                if len(grid) >= 2 or (len(grid) == 1 and len(grid[0]) >= 2):
                    formatted_text = format_table_to_dash_text(
                        grid,
                        separator=separator,
                        use_header=use_header,
                        show_row_indices=show_row_indices,
                        bullet_prefix=bullet_prefix
                    )
                    if formatted_text:
                        extracted_texts.append(formatted_text)
                        lines = formatted_text.splitlines()
                        first_p = consumed_paras[0]
                        first_p.text = lines[0]
                        first_p.paragraph_format.space_before = Pt(0)
                        first_p.paragraph_format.space_after = Pt(2)

                        parent_elm = first_p._element.getparent()
                        p_idx = parent_elm.index(first_p._element)

                        for l_offset, line_str in enumerate(lines[1:], 1):
                            np = doc.add_paragraph()
                            np.text = line_str
                            np.paragraph_format.space_before = Pt(0)
                            np.paragraph_format.space_after = Pt(2)
                            parent_elm.insert(p_idx + l_offset, np._element)

                        for rem_p in consumed_paras[1:]:
                            rem_elm = rem_p._element
                            if rem_elm.getparent() is not None:
                                rem_elm.getparent().remove(rem_elm)
                continue

        i += 1

    return extracted_texts

def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """
    Converts a .docx file to .pdf using MS Word COM interface (or docx2pdf fallback).
    Handles file locking gracefully if PDF is open in another viewer.
    """
    abs_docx = os.path.abspath(docx_path)
    abs_pdf = os.path.abspath(pdf_path)

    # If output PDF file exists and is locked, try removing or creating backup name
    if os.path.exists(abs_pdf):
        try:
            os.remove(abs_pdf)
        except Exception:
            base, ext = os.path.splitext(abs_pdf)
            abs_pdf = f"{base}_new{ext}"

    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(abs_docx)
            doc.SaveAs2(abs_pdf, FileFormat=17)  # 17 = wdFormatPDF
            doc.Close()
        finally:
            word.Quit()
    except Exception as e:
        try:
            from docx2pdf import convert
            convert(abs_docx, abs_pdf)
        except Exception:
            raise RuntimeError(f"Không thể xuất file PDF (Hãy đảm bảo file PDF không bị mở khoá trong phần mềm khác): {e}")


def process_document(file_path: str, separator: str = " | ", use_header: bool = True, show_row_indices: bool = False, bullet_prefix: bool = True, process_pseudo: bool = False) -> tuple[str, docx.Document]:
    """
    Reads doc/docx/pdf file, modifies ALL actual data tables IN-PLACE into compact bullet text lines.
    Preserves 100% of non-table text/headings/paragraphs intact.
    Returns (full_text_string, doc_out_object).
    """
    abs_input_path = os.path.abspath(file_path)
    real_docx_path, is_temp = convert_doc_or_pdf_to_docx(abs_input_path)

    try:
        doc = docx.Document(real_docx_path)
        extracted_texts = []

        # 1. Fix Table of Contents (Mục lục) concatenated paragraphs
        fix_toc_paragraphs(doc)

        tables_to_process = list(doc.tables)
        for table in tables_to_process:
            grid = clean_table_grid(table)
            
            # Case 1: Empty or pure header/footer boxes -> Remove without adding bullet text
            if is_header_footer_table(grid):
                tbl_elm = table._element
                if tbl_elm.getparent() is not None:
                    tbl_elm.getparent().remove(tbl_elm)
                continue
                
            num_rows = len(grid)
            num_cols = max(len(r) for r in grid) if num_rows > 0 else 0

            # Case 2: 1-column callout boxes or single-row list markers -> Unpack as clean normal text
            if num_cols == 1 or (num_rows == 1 and is_bullet_marker(grid[0][0])):
                unpack_text_table_inplace(doc, table, grid)
                continue

            # Case 3: Real Data Tables -> Convert to bullet key-value text lines
            table_text = format_table_to_dash_text(grid, separator=separator, use_header=use_header, show_row_indices=show_row_indices, bullet_prefix=bullet_prefix)
            if table_text:
                replace_table_inplace(doc, table, table_text)
                extracted_texts.append(table_text)

        # 2. Process Tabbed Pseudo-Tables (Only if explicitly enabled by user)
        if process_pseudo:
            pseudo_texts = process_pseudo_tables(doc, separator=separator, use_header=use_header, show_row_indices=show_row_indices, bullet_prefix=bullet_prefix)
            extracted_texts.extend(pseudo_texts)

        # 3. Clean up document spacing gaps completely
        cleanup_document_blank_spaces(doc)

        full_text = "\n\n".join(extracted_texts)
        return full_text, doc

    finally:
        if is_temp and os.path.exists(real_docx_path):
            try:
                os.remove(real_docx_path)
            except Exception:
                pass


def convert_file(file_path: str, output_path: str = None, export_txt: bool = False, export_pdf: bool = False, separator: str = " | ", use_header: bool = True, show_row_indices: bool = False, bullet_prefix: bool = True, process_pseudo: bool = False) -> str:
    """
    Main entry function to convert doc/docx/pdf file with IN-PLACE table replacement into bullet text.
    Target ONLY tables and preserve all non-table text intact.
    """
    full_text, doc_out = process_document(file_path, separator=separator, use_header=use_header, show_row_indices=show_row_indices, bullet_prefix=bullet_prefix, process_pseudo=process_pseudo)
    
    if output_path:
        out_lower = output_path.lower()
        if export_txt or out_lower.endswith(".txt"):
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(full_text)
        elif export_pdf or out_lower.endswith(".pdf"):
            temp_dir = tempfile.gettempdir()
            temp_docx = os.path.join(temp_dir, f"out_{os.path.splitext(os.path.basename(file_path))[0]}.docx")
            doc_out.save(temp_docx)
            try:
                convert_docx_to_pdf(temp_docx, output_path)
            finally:
                if os.path.exists(temp_docx):
                    try:
                        os.remove(temp_docx)
                    except Exception:
                        pass
        else:
            doc_out.save(output_path)
            
    return full_text

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Chuyển đổi BẢNG trong file PDF/Word thành text dạng gạch đầu dòng (In-place, chỉ làm việc với table)")
    parser.add_argument("input", help="Đường dẫn file PDF (.pdf) hoặc Word (.doc, .docx)")
    parser.add_argument("-o", "--output", help="Đường dẫn file đầu ra (.pdf, .docx hoặc .txt)")
    parser.add_argument("-s", "--separator", default=" | ", help="Ký tự phân cách (Mặc định: ' | ')")
    parser.add_argument("--txt", action="store_true", help="Xuất ra file text (.txt)")
    parser.add_argument("--pdf", action="store_true", help="Xuất ra file PDF (.pdf)")
    parser.add_argument("--no-header", action="store_true", help="Không sử dụng hàng đầu tiên làm tiêu đề")
    parser.add_argument("--indices", action="store_true", help="Hiển thị chỉ số dòng --- Dòng X ---")
    parser.add_argument("--no-bullet", action="store_true", help="Không dùng gạch đầu dòng (- )")
    parser.add_argument("--pseudo", action="store_true", help="Xử lý cả bảng giả lập dạng tab (mặc định tắt để tránh đụng tới văn bản khác)")

    args = parser.parse_args()
    try:
        res = convert_file(
            file_path=args.input,
            output_path=args.output,
            export_txt=args.txt,
            export_pdf=args.pdf,
            separator=args.separator,
            use_header=not args.no_header,
            show_row_indices=args.indices,
            bullet_prefix=not args.no_bullet,
            process_pseudo=args.pseudo
        )
        print("=== CHUYỂN ĐỔI BẢNG THÀNH CÔNG (CHỈ ĐỦNG VÀO TABLE) ===")
        print(res[:1000] if len(res) > 1000 else res)
    except Exception as err:
        print(f"Lỗi: {err}", file=sys.stderr)


