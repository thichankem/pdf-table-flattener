import sys
import os
import docx
from doc_table_converter import convert_doc_or_pdf_to_docx, clean_table_grid, format_table_to_dash_text, replace_table_inplace, clean_text_string

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

pdf_path = "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf"
real_docx_path, is_temp = convert_doc_or_pdf_to_docx(pdf_path)
doc = docx.Document(real_docx_path)

print(f"Raw doc tables count: {len(doc.tables)}")

def process_nested_tables_in_cell(cell):
    if not hasattr(cell, 'tables') or not cell.tables:
        return
    for nested_table in list(cell.tables):
        for nr in nested_table.rows:
            for nc in nr.cells:
                process_nested_tables_in_cell(nc)
        
        ngrid = clean_table_grid(nested_table)
        formatted_text = format_table_to_dash_text(ngrid)
        
        tbl_elm = nested_table._element
        parent_elm = tbl_elm.getparent()
        tbl_index = parent_elm.index(tbl_elm)
        
        if formatted_text:
            p = cell.add_paragraph()
            p.text = formatted_text
            parent_elm.insert(tbl_index, p._element)
        parent_elm.remove(tbl_elm)

for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            process_nested_tables_in_cell(c)

print("\n--- CHECKING TABLES AFTER NESTED PROCESSING ---")
for t_idx, t in enumerate(doc.tables):
    grid = clean_table_grid(t)
    print(f"\nTable {t_idx+1}: {len(grid)} rows x {max(len(r) for r in grid) if grid else 0} cols")
    for r_idx, r in enumerate(grid):
        r_str = " | ".join(r)
        if len(r_str) > 150:
            r_str = r_str[:150] + "..."
        print(f"  Row {r_idx}: {r_str}")
