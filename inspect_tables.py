import os
import sys
import docx

output_dir = "test_outputs"
pdf_files = [
    "Bảo hiểm nhân thọ - Lộc Phát Hưng Thịnh - Quy định sản phẩm.pdf",
    "Bảo hiểm nhân thọ - Lộc Phát Tràng An - Quy định sản phẩm.pdf",
    "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf",
    "Sản phẩm cho vay kinh doanh xi măng Xuân Thành - Kênh quầy.docx.pdf",
    "Sản phẩm cho vay tái tài trợ - kênh quầy.docx.pdf"
]

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

for idx, pdf_name in enumerate(pdf_files, 1):
    out_docx = os.path.join(output_dir, f"out_{idx}.docx")
    print(f"\n==================================================")
    print(f"FILE {idx}: {pdf_name}")
    print(f"==================================================")
    if not os.path.exists(out_docx):
        print("File out docx not found")
        continue
    doc = docx.Document(out_docx)
    
    table_lines = []
    normal_lines = []
    empty_lines = []
    
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if not t.strip():
            empty_lines.append(i)
        elif t.strip().startswith("- "):
            table_lines.append((i, t.strip()))
        else:
            normal_lines.append((i, t.strip()))
            
    print(f"Tổng số dòng: {len(doc.paragraphs)}")
    print(f"Dòng trống: {len(empty_lines)}")
    print(f"Dòng bảng (- ...): {len(table_lines)}")
    
    print("\n--- TẤT CẢ CÁC DÒNG BẢNG CONVERT ĐƯỢC ---")
    for p_i, line in table_lines:
        print(f"[{p_i:03d}]: {line}")
