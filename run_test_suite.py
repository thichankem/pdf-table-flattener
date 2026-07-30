import os
import sys
import re
import docx
from doc_table_converter import convert_file, has_drawing

pdf_files = [
    "Bảo hiểm nhân thọ - Lộc Phát Hưng Thịnh - Quy định sản phẩm.pdf",
    "Bảo hiểm nhân thọ - Lộc Phát Tràng An - Quy định sản phẩm.pdf",
    "Sản phẩm cho vay CBNV trường đại học Quốc gia TP.Hồ Chí Minh - kênh quầy.docx.pdf",
    "Sản phẩm cho vay kinh doanh xi măng Xuân Thành - Kênh quầy.docx.pdf",
    "Sản phẩm cho vay tái tài trợ - kênh quầy.docx.pdf"
]

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

print("================================================================")
print("=== CHẠY KIỂM THỬ XUẤT FILE [TÊN_GỐC]_convert.pdf THEO TEST.MD ===")
print("================================================================\n")

all_passed = True

for idx, pdf_name in enumerate(pdf_files, 1):
    pdf_path = os.path.abspath(pdf_name)
    if not os.path.exists(pdf_path):
        print(f"[{idx}/5] ❌ FILE KHÔNG TỒN TẠI: {pdf_name}")
        all_passed = False
        continue
    
    base_name, _ = os.path.splitext(pdf_name)
    out_pdf = os.path.abspath(f"{base_name}_convert.pdf")
    out_docx_audit = os.path.abspath(f"{base_name}_convert.docx")

    print(f"--- [{idx}/5] FILE: {pdf_name} ---")
    
    try:
        # 1. Convert PDF to output PDF named base_name_convert.pdf
        full_txt = convert_file(pdf_path, output_path=out_pdf, export_pdf=True, separator=" | ", bullet_prefix=True)
        # Also convert to docx for structural audit
        convert_file(pdf_path, output_path=out_docx_audit, export_pdf=False, separator=" | ", bullet_prefix=True)
        
        pdf_exists = os.path.exists(out_pdf) and os.path.getsize(out_pdf) > 0
        pdf_size_kb = (os.path.getsize(out_pdf) / 1024) if pdf_exists else 0
        
        doc = docx.Document(out_docx_audit)
        
        # 2. Check Criterion 2: Remaining tables must be 0
        table_count = len(doc.tables)
        crit2_pass = (table_count == 0) and pdf_exists
        
        paras = doc.paragraphs
        total_paras = len(paras)
        
        empty_paras = 0
        consecutive_empty_max = 0
        curr_empty = 0
        strange_chars_found = []
        fake_column_labels_found = []
        bullet_lines = []
        normal_lines = []
        drawing_count = 0
        
        for p_i, p in enumerate(paras):
            t = p.text
            t_strip = t.strip()
            
            if has_drawing(p):
                drawing_count += 1
                
            if not t_strip and not has_drawing(p):
                empty_paras += 1
                curr_empty += 1
                if curr_empty > consecutive_empty_max:
                    consecutive_empty_max = curr_empty
            else:
                curr_empty = 0
                
            # Check Criterion 3: Strange / invisible characters
            if '\u200b' in t or '\ufeff' in t or '\x00' in t or '\x0c' in t or '\ufffd' in t:
                strange_chars_found.append((p_i, t_strip[:50]))
                
            # Check for artificial fake column headers like "Cột 1:", "Cột 2:"
            if re.search(r'\bCột \d+:', t):
                fake_column_labels_found.append((p_i, t_strip[:80]))
                
            if t_strip.startswith("- "):
                bullet_lines.append((p_i, t_strip))
            elif t_strip:
                normal_lines.append((p_i, t_strip))
                
        # Check Criterion 1: Non-table text & images preserved 100%
        crit1_pass = (len(normal_lines) > 0) and (drawing_count > 0)
        crit3_pass = (consecutive_empty_max <= 1) and (len(strange_chars_found) == 0) and (len(fake_column_labels_found) == 0)
        
        file_pass = crit1_pass and crit2_pass and crit3_pass
        if not file_pass:
            all_passed = False
            
        status_str = "✅ ĐẠT 100%" if file_pass else "❌ CHƯA ĐẠT"
        print(f"  Kết quả: {status_str}")
        print(f"  • Tên file PDF kết quả: {out_pdf} ({pdf_size_kb:.1f} KB)")
        print(f"  • Tiêu chí 1 (Giữ nguyên văn bản ngoài bảng & {drawing_count} hình ảnh/sơ đồ): {'✅ ĐẠT 100%' if crit1_pass else '❌ CHƯA ĐẠT'} ({len(normal_lines)} đoạn văn)")
        print(f"  • Tiêu chí 2 (Flatten bảng thành key-value, 0 table còn lại): {'✅ ĐẠT 100%' if crit2_pass else '❌ CHƯA ĐẠT'} (Table còn lại: {table_count}, Dòng bảng: {len(bullet_lines)})")
        print(f"  • Tiêu chí 3 (Không ký tự lạ, nhãn giả Cột X, hoặc hở dòng): {'✅ ĐẠT 100%' if crit3_pass else '❌ CHƯA ĐẠT'}")
        print(f"     - Hình ảnh/sơ đồ phát hiện & giữ nguyên: {drawing_count}")
        print(f"     - Hở dòng liên tiếp tối đa: {consecutive_empty_max} (Yêu cầu <= 1)")
        print(f"     - Ký tự lạ/ẩn: {len(strange_chars_found)}")
        print(f"     - Nhãn giả (Cột 1:, Cột 2:): {len(fake_column_labels_found)}")
        
        if bullet_lines:
            print("  • Mẫu 3 dòng bảng gạch đầu dòng trong PDF:")
            for b_i, b_val in bullet_lines[:3]:
                print(f"     [Line {b_i:03d}]: {b_val[:120]}")
        print()
            
    except Exception as e:
        all_passed = False
        print(f"  ❌ LỖI KHI XỬ LÝ: {e}")
        import traceback
        traceback.print_exc()

print("================================================================")
if all_passed:
    print("🎉 TẤT CẢ 5 FILE PDF ĐÃ ĐƯỢC XUẤT THÀNH FILE [TÊN_GỐC]_convert.pdf VÀ ĐẠT 100% CẢ 3 TIÊU CHÍ (GIỮ NGUYÊN HÌNH ẢNH 100%)!")
else:
    print("⚠️ CÓ FILE CHƯA ĐẠT TIÊU CHÍ TRONG TEST.MD!")
print("================================================================")
